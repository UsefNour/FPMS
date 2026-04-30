"""Generate completed FET Ethical Review Checklist for FPMS."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

BLACK = RGBColor(0x00,0x00,0x00)
GREY  = RGBColor(0x59,0x59,0x59)
ORANGE = RGBColor(0xED,0x7D,0x31)

def para(text='', size=11, bold=False, italic=False, color=BLACK,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size); r.bold = bold; r.italic = italic
        r.font.color.rgb = color
    return p

def cell_set(cell, text, bold=False, size=10.5, color=BLACK, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, shade=None):
    if shade:
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), shade); tcPr.append(shd)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    r.font.color.rgb = color; r.italic = italic

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

# ── Header ────────────────────────────────────────────────────────────────────
para('Faculty of Environment & Technology', size=12, bold=True,
     align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
para('Faculty Research Ethics Committee (FREC)', size=11,
     align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

para('Ethical Review Checklist for Undergraduate and Postgraduate Modules',
     size=14, bold=True, space_before=4, space_after=6)

para('Staff and PG research students must not use this form, but should instead, if appropriate, '
     'submit a full application for ethical approval to the Faculty Research Ethics Committee (FREC).',
     size=10.5, space_after=4)
para('Please provide project details and complete the checklist below.',
     size=10.5, italic=True, space_after=10)

# ── Project Details ───────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('Project Details:'); r.bold = True; r.font.size = Pt(11)
p.paragraph_format.space_after = Pt(4)

t1 = doc.add_table(rows=5, cols=2)
t1.style = 'Table Grid'
set_col_width(t1, 0, 5.0); set_col_width(t1, 1, 11.5)
rows1 = [
    ('Module name',          'Software Development Project'),
    ('Module code',          'UFCFFF-30-3'),
    ('Module leader',        'Steve Battle'),
    ('Project Supervisor',   'Steve Battle'),
    ('Proposed project title','Fighter Performance Management System (FPMS)'),
]
for i,(label,val) in enumerate(rows1):
    cell_set(t1.rows[i].cells[0], label, bold=True)
    cell_set(t1.rows[i].cells[1], val)

doc.add_paragraph()

# ── Applicant Details ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('Applicant Details:'); r.bold = True; r.font.size = Pt(11)
p.paragraph_format.space_after = Pt(4)

t2 = doc.add_table(rows=3, cols=2)
t2.style = 'Table Grid'
set_col_width(t2, 0, 5.0); set_col_width(t2, 1, 11.5)
rows2 = [
    ('Name of Student',        'Youssef Nour'),
    ('Student Number',         '23019868'),
    ("Student's email address",'youssef2.nour@live.uwe.ac.uk'),
]
for i,(label,val) in enumerate(rows2):
    cell_set(t2.rows[i].cells[0], label, bold=True)
    cell_set(t2.rows[i].cells[1], val)

doc.add_paragraph()

# ── Checklist Table ───────────────────────────────────────────────────────────
questions = [
    (
        '1.',
        'Does the proposed project involve human tissue, human participants, '
        'animals, environmental damage, or the NHS.',
        'No',
        'FPMS is a purely software-based development project. No human participants '
        'were recruited, no tissue collected, no animals involved, and no NHS '
        'engagement. All testing used synthetic data created by the developer.'
    ),
    (
        '2.',
        'Will participants be clearly asked to give consent to take part in the '
        'research and informed about how data collected in the research will be used?',
        'N/A',
        'No human participants are involved in this project.'
    ),
    (
        '3.',
        'If they choose, can a participant withdraw at any time (prior to a point '
        'of "no return" in the use of their data)? Are they told this?',
        'N/A',
        'No human participants are involved in this project.'
    ),
    (
        '4.',
        'Are measures in place to provide confidentiality for participants and '
        'ensure secure management and disposal of data collected from them?',
        'Yes',
        'The application stores user data securely: passwords are hashed with bcrypt '
        '(Werkzeug), all forms carry CSRF protection (Flask-WTF), and sessions are '
        'managed via Flask-Login. The test database (fpms.db) is held locally, '
        'excluded from the public GitHub repository via .gitignore, and will be '
        'deleted after assessment.'
    ),
    (
        '5.',
        'Does the study involve people who are particularly vulnerable or unable to '
        'give informed consent (eg, children or people with learning difficulties)?',
        'No',
        ''
    ),
    (
        '6.',
        'Could your research cause stress, physical or psychological harm to '
        'humans or animals, or environmental damage?',
        'No',
        ''
    ),
    (
        '7.',
        'Could any aspects of the research lead to unethical behaviour by '
        'participants or researchers (eg, invasion of privacy, deceit, coercion, '
        'fraud, abuse)?',
        'No',
        ''
    ),
    (
        '8.',
        'Does the research involve the NHS or collection or storage of human tissue '
        '(includes anything containing human cells, such as saliva and urine)?',
        'No',
        ''
    ),
]

tq = doc.add_table(rows=len(questions)+1, cols=4)
tq.style = 'Table Grid'
set_col_width(tq, 0, 0.6)
set_col_width(tq, 1, 8.4)
set_col_width(tq, 2, 1.5)
set_col_width(tq, 3, 6.0)

# Header row
for col, hdr in enumerate(['', 'CHECKLIST QUESTIONS', 'Yes/No', 'Explanation']):
    cell_set(tq.rows[0].cells[col], hdr, bold=True, size=10.5,
             shade='D9D9D9' if hdr else 'D9D9D9')

for i, (num, question, ans, explanation) in enumerate(questions):
    row = tq.rows[i+1]
    cell_set(row.cells[0], num, bold=True, size=10.5)

    # Question text — orange for Q2-4 (as in the example form)
    c = ORANGE if num in ('2.','3.','4.') else BLACK
    cell_set(row.cells[1], question, size=10.5, color=c)

    # Yes/No cell
    cell_set(row.cells[2], ans, size=10.5,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    # Explanation
    cell_set(row.cells[3], explanation, size=10.5)

doc.add_paragraph()

# ── Risk guidance text ────────────────────────────────────────────────────────
guidance = doc.add_paragraph()
guidance.paragraph_format.space_after = Pt(4)
r = guidance.add_run(
    'Your explanations should indicate briefly for Qs 2-4 how these requirements '
    'will be met, and for Qs 5-8 what the pertinent concerns are.')
r.font.size = Pt(10.5)

def bullet(text, bold_part='', rest=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_part:
        rb = p.add_run(bold_part); rb.bold = True; rb.font.size = Pt(10.5)
        rr = p.add_run(rest); rr.font.size = Pt(10.5)
    else:
        r = p.add_run(text); r.font.size = Pt(10.5)

bullet('', 'Minimal Risk: ', 'If Q 1 is answered ‘No’, then no ethics approval is needed.')
bullet('', 'Low Risk: ',
       'If Qs 2-4 are answered ‘Yes’ and Qs 5-8 are answered ‘No’, then no approval is '
       'needed from the Faculty Research Ethics Committee (FREC). However, your supervisor must '
       'approve (a) your information and consent forms (Qs 2 & 3) and (b) your measures for '
       'participant confidentiality and secure data management (Q4).')
bullet('', 'High Risk: ',
       'If any of Qs 5-8 are answered ‘Yes’, then you must submit an application for full '
       'ethics approval before the project can start. This can take up to 6 weeks. Consult your '
       'supervisor about how to apply for full ethics approval.')

doc.add_paragraph()
p = doc.add_paragraph()
rb = p.add_run('Risk Assessment: '); rb.bold = True; rb.font.size = Pt(10.5)
rr = p.add_run(
    'Separate guidance on risk assessment can be found on UWE’s Health and Safety '
    'forms webpage at https://go.uwe.ac.uk/RiskAssessment. If needed, you must complete a Risk '
    'Assessment form. This must also be attached to your application for full ethics approval if '
    'your project is High Risk.')
rr.font.size = Pt(10.5)
p.paragraph_format.space_after = Pt(10)

# ── Supervisor / submission boxes ─────────────────────────────────────────────
for box_text in [
    'Your supervisor must check your responses above before you submit this form.',
    'Submit this completed form via the Assignments area in Blackboard (or elsewhere if so directed by the module leader or your supervisor).',
    'After you have uploaded this form, your supervisor will confirm it has been correctly completed by “marking” it as Passed/100% via the My Grades link on the Blackboard.',
]:
    tb = doc.add_table(rows=1, cols=1)
    tb.style = 'Table Grid'
    p = tb.rows[0].cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(box_text); r.font.size = Pt(10.5)
    if box_text.startswith('Your supervisor'):
        r.bold = True
    doc.add_paragraph()

para('Further research ethics guidance is available at http://www1.uwe.ac.uk/research/researchethics',
     size=10.5, space_before=4)

# ── Footer note ───────────────────────────────────────────────────────────────
para('FET FREC – UG/PGR Ethical Review Module Checklist    v20 on 5 Aug 2016',
     size=9, color=GREY, space_before=20)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(sys.argv[0])),
                   'FPMS_Ethical_Review_Checklist.docx')
doc.save(out)
print(f'Saved: {out}')
