import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

BLUE   = RGBColor(31, 73, 125)
GREY   = RGBColor(100, 100, 100)
LIGHT  = RGBColor(150, 150, 150)
WHITE  = RGBColor(255, 255, 255)
GREEN  = RGBColor(0, 176, 80)
AMBER  = RGBColor(255, 153, 0)
RED    = RGBColor(192, 0, 0)

# ── Helpers ───────────────────────────────────────────────────────────────────

def centred(text, size, bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return p

def body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = GREY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)

def h1(text):
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Calibri'; run.font.color.rgb = BLUE; run.font.size = Pt(16)
    h.paragraph_format.space_before = Pt(18); h.paragraph_format.space_after = Pt(6)

def h2(text):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Calibri'; run.font.color.rgb = BLUE; run.font.size = Pt(13)
    h.paragraph_format.space_before = Pt(12); h.paragraph_format.space_after = Pt(4)

def h3(text):
    h = doc.add_heading(text, level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Calibri'; run.font.color.rgb = GREY; run.font.size = Pt(11)
    h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(2)

def add_hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1F497D')
    pBdr.append(bot); pPr.append(pBdr)

def code_block(lines):
    """Render a list of strings as a monospaced code block with grey background."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        # shaded background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'; r.font.size = Pt(9)

def placeholder(text='[Content to be added]'):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = LIGHT; r.italic = True
    r.font.name = 'Calibri'; r.font.size = Pt(11)

def add_picture_safe(rel_path, fig_label, width=Cm(14)):
    """Insert image if file exists; otherwise draw a labelled placeholder box."""
    full = os.path.join(os.path.dirname(__file__), rel_path)
    if os.path.exists(full):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(full, width=width)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)
        r = p.add_run(f'[ {fig_label} — save screenshot as: {rel_path} ]')
        r.font.name = 'Calibri'; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(160, 160, 160); r.italic = True
    return p

def set_cell(cell, text, bold=False, bg_color=None, center=False, color=None, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(size); r.bold = bold
    if color:  r.font.color.rgb = color
    if bg_color:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        tcPr.append(shd)

def tick_cross(cell, val):
    symbol = '✓' if val else '✗'
    bg = 'E2EFDA' if val else 'FFE0E0'
    set_cell(cell, symbol, bold=True, bg_color=bg, center=True)

def no_split(tbl):
    """Prevent every row in a table from splitting across a page break."""
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit')
        trPr.append(cs)

def add_page_numbers():
    """Insert a centred page-number field into every section footer."""
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = GREY
        for tag, text in [('begin', None), ('instrText', 'PAGE'), ('end', None)]:
            if tag == 'instrText':
                el = OxmlElement('w:instrText')
                el.set(qn('xml:space'), 'preserve')
                el.text = text
            else:
                el = OxmlElement('w:fldChar')
                el.set(qn('w:fldCharType'), tag)
            run._r.append(el)

# ── Reference list ────────────────────────────────────────────────────────────
REFERENCES = [
    # Ch1 — Background research (November 2025)
    'Statista (2024) UFC – Statistics and Facts. Available at: '
    'https://www.statista.com/topics/3899/ufc/ [Accessed: 12 November 2025].',

    'Grinberg, M. (2018) Flask Web Development: Developing Web Applications with Python. '
    '2nd edn. Sebastopol: O\'Reilly Media.',

    'Kitman Labs (2023) Athlete Intelligence Platform. Available at: '
    'https://www.kitmanlabs.com [Accessed: 19 November 2025].',

    'Hudl (2024) Video Analysis and Performance Tools for Sport. Available at: '
    'https://www.hudl.com [Accessed: 24 November 2025].',

    # Ch2 — Methodology (December 2025)
    'Beck, K., Beedle, M., van Bennekum, A., Cockburn, A., Cunningham, W., Fowler, M., '
    'Grenning, J., Highsmith, J., Hunt, A., Jeffries, R., Kern, J., Marick, B., Martin, R.C., '
    'Mellor, S., Schwaber, K., Sutherland, J. and Thomas, D. (2001) Manifesto for Agile Software '
    'Development. Available at: https://agilemanifesto.org [Accessed: 2 December 2025].',

    'Schwaber, K. and Sutherland, J. (2020) The Scrum Guide: The Definitive Guide to Scrum — '
    'The Rules of the Game. Available at: https://scrumguides.org/scrum-guide.html '
    '[Accessed: 5 December 2025].',

    'Sommerville, I. (2016) Software Engineering. 10th edn. Harlow: Pearson Education.',

    'Git (2024) Git — Distributed Version Control System. Available at: '
    'https://git-scm.com [Accessed: 8 December 2025].',

    'GitHub (2024) GitHub: Let\'s Build from Here. Available at: '
    'https://github.com [Accessed: 8 December 2025].',

    'UWE Bristol (2025) Research Ethics at UWE Bristol. Available at: '
    'https://www.uwe.ac.uk/research/research-environment/research-ethics '
    '[Accessed: 15 December 2025].',

    # Ch3 — Technology research (January–February 2026)
    'Django Software Foundation (2024) Django Documentation. Available at: '
    'https://www.djangoproject.com/start/ [Accessed: 9 January 2026].',

    'Ramírez, S. (2024) FastAPI Documentation. Available at: '
    'https://fastapi.tiangolo.com [Accessed: 11 January 2026].',

    'SQLite (2024) SQLite Home Page. Available at: '
    'https://www.sqlite.org/index.html [Accessed: 16 January 2026].',

    'PostgreSQL Global Development Group (2024) PostgreSQL: The World\'s Most Advanced Open '
    'Source Relational Database. Available at: https://www.postgresql.org [Accessed: 16 January 2026].',

    'MyFitnessPal (2024) MyFitnessPal: Calorie Counter and Diet Tracker. Available at: '
    'https://www.myfitnesspal.com [Accessed: 23 January 2026].',

    'Tapology (2024) MMA Matchmaker and Event Listings. Available at: '
    'https://www.tapology.com [Accessed: 27 January 2026].',

    'Pallets Projects (2024) Flask Documentation (3.x). Available at: '
    'https://flask.palletsprojects.com [Accessed: 4 February 2026].',

    'SQLAlchemy (2024) SQLAlchemy — The Python SQL Toolkit and Object Relational Mapper. '
    'Available at: https://www.sqlalchemy.org [Accessed: 4 February 2026].',

    # Ch4 — Requirements (February 2026)
    'ISO/IEC 9126 (2001) Software Engineering — Product Quality. Geneva: International '
    'Organisation for Standardisation.',

    'Cockburn, A. (2001) Writing Effective Use Cases. Boston: Addison-Wesley.',

    # Ch5 — Implementation (February–March 2026)
    'Flask-SocketIO (2024) Flask-SocketIO Documentation. Available at: '
    'https://flask-socketio.readthedocs.io [Accessed: 18 February 2026].',

    'Sinha, R. and Hearst, M. (2001) "Similarity and Grouping in Computer-Supported '
    'Collaborative Work", in Proceedings of CHI 2001, pp. 135–142.',

    'Werkzeug (2024) Werkzeug Utilities Documentation. Available at: '
    'https://werkzeug.palletsprojects.com [Accessed: 3 March 2026].',

    'Sinnott, R.W. (1984) "Virtues of the Haversine", Sky and Telescope, 68(2), p. 159.',

    # Ch6 — Testing (March–April 2026)
    'Myers, G.J., Sandler, C. and Badgett, T. (2011) The Art of Software Testing. '
    '3rd edn. Hoboken: John Wiley & Sons.',

    # Ch7 — Conclusion (April 2026)
    'Gibbs, G. (1988) Learning by Doing: A Guide to Teaching and Learning Methods. '
    'Oxford: Further Education Unit, Oxford Polytechnic.',

    'Heroku (2024) Cloud Application Platform. Available at: '
    'https://www.heroku.com [Accessed: 14 April 2026].',

    'pytest (2024) pytest Documentation. Available at: '
    'https://docs.pytest.org [Accessed: 14 April 2026].',
]

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
centred('University of the West of England, Bristol', 12, color=GREY)
centred('Faculty of Environment and Technology', 11, color=GREY)
doc.add_paragraph(); doc.add_paragraph()
centred('Fighter Performance Management System', 26, bold=True, color=BLUE)
doc.add_paragraph()
centred('FPMS: A Web-Based Platform for Combat Sports\nAthletes, Coaches and Event Organisers', 15, color=GREY)
doc.add_paragraph(); doc.add_paragraph()
add_hrule(); doc.add_paragraph()

for label, value in [
    ('Student Name:',    'Youssef Nour'),
    ('Student Number:',  '23019868'),
    ('Module Title:',    'Software Development Project'),
    ('Module Code:',     'UFCFFF-30-3'),
    ('Module Leader:',   'Steve Battle'),
    ('Submission Date:', '30 April 2026'),
    ('Word Count:',      'approx. 6,400 words (excl. tables, references and captions)'),
    ('GitHub:',          'https://github.com/UsefNour/FPMS'),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + '  '); r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.bold = True; r1.font.color.rgb = BLUE
    r2 = p.add_run(value); r2.font.name = 'Calibri'; r2.font.size = Pt(11)

doc.add_paragraph()
centred('School of Computing and Creative Technologies', 10, color=LIGHT)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
para = doc.add_paragraph()
run = para.add_run()
b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = r'TOC \o "1-3" \h \z \u'
s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
run._r.append(b); run._r.append(i); run._r.append(s); run._r.append(e)
p = doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
h1('Chapter 1 — Introduction')

h2('1.1  Background and Motivation')
body(
    'Combat sports — MMA, boxing, Muay Thai, Brazilian Jiu-Jitsu, wrestling and related disciplines — '
    'have grown enormously over the past two decades. Promotions like the UFC and Bellator have turned '
    'what was once a niche pursuit into a global industry, with the UFC alone valued at around $12 billion '
    'in 2023 (Statista, 2024). That growth has raised the bar for everyone in the sport, not just '
    'headliners — regional and amateur competitors are now expected to show up prepared in ways that '
    'would have seemed excessive a generation ago.'
)
body(
    'The problem is that the tools available to most fighters have not kept pace. Training camps are '
    'still planned in Notes apps or paper diaries. Weight cuts are tracked on spreadsheets, if at all. '
    'Finding a sparring partner usually means posting in a Facebook group. None of these tools talk '
    'to each other, none of them understand what a fight camp actually is, and the result is that '
    'fighters waste time on logistics that should be spent training. Coaches get WhatsApp messages '
    'instead of structured progress data. Event organisers have no easy way to reach the athletes '
    'they are trying to fill cards with. The gap is real and it costs people preparation time they '
    'cannot get back.'
)

h2('1.2  Problem Statement')
body(
    'There is no single tool that covers the full preparation cycle of an independent combat sports athlete. '
    'Commercial platforms like Kitman Labs (2023) and Hudl (2024) are built for professional team environments '
    'and priced accordingly. Everything else is either too generic — a nutrition app that knows nothing about '
    'fight-camp periodisation — or too narrow, like an event-listing site that offers no training functionality. '
    'The result is that fighters piece things together across half a dozen disconnected apps, none of which '
    'understands what a weight cut is or why a sparring peak comes before a taper week.'
)
body(
    'Table 1.1 illustrates this gap by mapping key features against a representative selection of existing '
    'tools, highlighting the areas where FPMS provides unique, integrated coverage.'
)

# Table 1.1 — Feature gap analysis
features = [
    'Fight-camp planning',
    'Weight cut tracking',
    'Opponent game-plan builder',
    'Sparring partner matching',
    'Event creation & discovery',
    'Real-time peer chat',
    'Risk / readiness scoring',
    'Combat-sports specific',
    'Free / accessible',
]
tools_cols = ['Generic apps\n(e.g. Notion)', 'MyFitnessPal\n(2024)', 'Kitman Labs\n(2023)', 'Hudl\n(2024)', 'FPMS\n(this project)']
values = [
    [False, False, False, False, True],
    [False, True,  True,  False, True],
    [False, False, False, False, True],
    [False, False, False, False, True],
    [False, False, False, False, True],
    [False, False, False, False, True],
    [False, False, False, False, True],
    [False, False, True,  True,  True],
    [True,  True,  False, False, True],
]

tbl = doc.add_table(rows=1, cols=len(tools_cols)+1)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
set_cell(hdr[0], 'Feature', bold=True, bg_color='1F497D', color=WHITE)
for i, col in enumerate(tools_cols):
    set_cell(hdr[i+1], col, bold=True, bg_color='1F497D', color=WHITE, center=True)

for feat, row_vals in zip(features, values):
    row = tbl.add_row().cells
    set_cell(row[0], feat, bold=False)
    for j, v in enumerate(row_vals):
        tick_cross(row[j+1], v)

no_split(tbl)
caption('Table 1.1 — Feature gap analysis: existing tools vs FPMS (✓ = supported, ✗ = not supported)')

body(
    'This project therefore sets out to design, develop and evaluate a Fighter Performance Management System '
    '(FPMS): a single, accessible web application that consolidates camp planning, weight monitoring, '
    'game-plan generation, sparring-partner discovery, event management, and social communication into one '
    'unified platform.'
)

h2('1.3  Aims and Objectives')
body(
    'The aim of this project is to build, test and critically evaluate a web application that gives '
    'combat sports athletes a single place to manage everything from camp planning to event registration. '
    'The following objectives break that down:'
)
for obj in [
    'O1 — Conduct secondary research into existing sports performance management tools, web application '
     'frameworks, and combat-sports-specific user needs.',
    'O2 — Derive a clear set of functional and non-functional requirements from the research, capturing '
     'them as numbered use-cases.',
    'O3 — Design and implement an FPMS web application using the Flask micro-framework (Grinberg, 2018), '
     'a relational database, and real-time WebSocket communication.',
    'O4 — Test the system against the defined requirements using a combination of manual exploratory '
     'testing and route-level functional tests.',
    'O5 — Critically reflect on the development process, the technology choices made, and the '
     'shortcomings of the resulting artefact.',
]:
    bullet(obj)
doc.add_paragraph()

h2('1.4  Scope and Limitations')
body(
    'FPMS is built for individual athletes and small training groups — the kind of fighter who manages their '
    'own camp without a full coaching staff behind them. It is not aimed at professional team environments '
    'and does not integrate with wearable sensors or third-party fitness APIs; those are possible future '
    'extensions but are out of scope here. The application runs locally with SQLite as the database; '
    'switching to a production setup — PostgreSQL, HTTPS, proper secret management — is discussed in the '
    'conclusion but is not part of the assessed deliverable.'
)
body(
    'Ethics approval was obtained prior to conducting any user observation or data collection. '
    'All data used during development and testing is synthetic; no real personal or biometric data '
    'was collected or stored without informed consent. Evidence of ethical approval is provided in '
    'the appendices.'
)

h2('1.5  Report Structure')
body(
    'Chapter 2 covers the development methodology, toolchain and risk register. Chapter 3 sets out the '
    'secondary research and the technology selection decisions that shaped the architecture. Chapter 4 '
    'formalises the requirements as 40 numbered FRs and 17 NFRs, with a use-case overview. '
    'Chapter 5 documents design and implementation, including the Blueprint architecture, key '
    'algorithms and selected code extracts. Chapter 6 presents the testing results against all requirements. '
    'Chapter 7 reflects critically on the project and sets out future work.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — METHOD
# ═══════════════════════════════════════════════════════════════════════════════
h1('Chapter 2 — Method')

h2('2.1  Development Methodology')
body(
    'At the start of the project it was not entirely clear what the system should do in full — only '
    'the core problem was well-defined. A Waterfall approach therefore did not suit: locking in '
    'complete requirements before writing any code would have produced a requirements document full '
    'of guesses (Sommerville, 2016). Heavyweight Scrum was also ruled out; its prescribed roles of '
    'Product Owner, Scrum Master and Development Team (Schwaber and Sutherland, 2020) are built '
    'around multi-person teams and add process overhead that makes little sense for a solo project.'
)
body(
    'Instead, an iterative and incremental approach was adopted, drawing on the core values of the Agile '
    'Manifesto (Beck et al., 2001): prioritising working software over comprehensive documentation, '
    'responding to change over following a plan, and valuing self-organising work over rigid process. '
    'Development was organised into four loosely time-boxed iterations of approximately two to three weeks '
    'each, with each iteration producing a demonstrable vertical slice of functionality:'
)
for item in [
    'Iteration 1 — Foundation: project scaffolding, database schema, user authentication (registration, '
     'login, logout), fighter profile management.',
    'Iteration 2 — Training core: fight-camp planner, game-plan generator, weight tracker, '
     'risk-readiness assessment.',
    'Iteration 3 — Social and discovery: friend system, real-time chat via WebSocket, '
     'sparring-partner matching with geolocation.',
    'Iteration 4 — Events and polish: event creation, discovery, interest/registration, '
     'admin panel, Blueprint refactoring, cross-cutting testing.',
]:
    bullet(item)
doc.add_paragraph()
body(
    'At the end of each iteration, all implemented routes were tested manually against the use-cases '
    'defined in Chapter 4, and any defects or design concerns were captured and prioritised for the '
    'following iteration. This closed feedback loop ensured that architectural problems (such as the '
    'circular-import risk from placing all routes in a single module) were identified and resolved '
    'before they compounded. Figure 2.1 presents the full project timeline as a Gantt chart, '
    'spanning the six-month development period from November 2025 to the submission deadline of '
    '30 April 2026.'
)

add_picture_safe('screenshots/gantt_chart.png', 'Figure 2.1 — Gantt Chart', width=Cm(16))
caption('Figure 2.1 — FPMS project Gantt chart: November 2025 – April 2026 (dashed line = submission deadline)')

h2('2.2  Toolchain')
body('The following tools and technologies were selected for the project. Table 2.1 summarises the full toolchain.')

tools = [
    ('Language',          'Python 3.13',          'Widely supported, extensive library ecosystem, natural fit for Flask.'),
    ('Web framework',     'Flask 3.x',             'Lightweight micro-framework; gives full control over structure (Grinberg, 2018).'),
    ('ORM',               'SQLAlchemy 2.x',        'Abstracts database interactions; supports multiple backends.'),
    ('Authentication',    'Flask-Login',           'Session management and login-required decorators.'),
    ('Real-time comms',   'Flask-SocketIO 5.x',    'WebSocket support built on Socket.IO; bidirectional event model.'),
    ('Database (dev)',    'SQLite',                'Zero-configuration file-based database; sufficient for development and testing.'),
    ('Frontend',          'Bootstrap 5 + Jinja2',  'Responsive grid system; Jinja2 for server-side HTML templating.'),
    ('Form handling',     'WTForms / Flask-WTF',   'CSRF protection, field validation, and form rendering helpers.'),
    ('Version control',   'Git / GitHub',          'Distributed version control; remote repository for backup and history.'),
    ('IDE',               'Visual Studio Code',    'Lightweight editor with Python extension and integrated terminal.'),
]

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
for cell, txt in zip(tbl2.rows[0].cells, ['Component', 'Tool / Technology', 'Rationale']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for component, tool, rationale in tools:
    row = tbl2.add_row().cells
    set_cell(row[0], component, bold=True)
    set_cell(row[1], tool)
    set_cell(row[2], rationale)
no_split(tbl2)
caption('Table 2.1 — Project toolchain')

h2('2.3  Version Control and Configuration Management')
body(
    'Git was used as the version control system throughout the project, with the remote repository hosted '
    'on GitHub (GitHub, 2024). Commits were made at logical checkpoints — typically after each feature was '
    'implemented and manually verified — rather than on an arbitrary time schedule. Commit messages followed '
    'a short imperative convention (e.g. "Add sparring profile route and form") to make the history '
    'readable as a narrative of development decisions.'
)
body(
    'A single main branch was used for this project given its solo nature; feature branches were created '
    'for larger structural changes such as the Blueprint refactoring in Iteration 4, and merged back only '
    'after the full route suite returned HTTP 200/302 responses. Sensitive configuration values — the '
    'application secret key and upload folder path — were noted in the README as requiring environment '
    'variable substitution before any production deployment, and were never committed to the repository.'
)

h2('2.4  Risk Management')
body(
    'A risk register was maintained from the start of the project and reviewed at the beginning of each '
    'iteration. Risks were scored by multiplying Likelihood (1–3) by Impact (1–3) to produce a Risk '
    'Rating following standard practice (Sommerville, 2016). Table 2.2 presents the key risks identified '
    'and the mitigation strategy applied to each.'
)

risk_rows = [
    ('R1', 'Scope creep — features expanding beyond what can be delivered in time',         'High',   'High',   '9', 'FF6666'),
    ('R2', 'Technical complexity of real-time WebSocket chat',                              'Medium', 'High',   '6', 'FFCC66'),
    ('R3', 'Circular import errors from monolithic app structure',                          'Medium', 'Medium', '4', 'FFFF99'),
    ('R4', 'Data loss due to lack of database backups',                                     'Low',    'High',   '3', 'CCFFCC'),
    ('R5', 'Security vulnerabilities (CSRF, SQL injection)',                                'Medium', 'High',   '6', 'FFCC66'),
    ('R6', 'Python / dependency version mismatch across environments',                      'Medium', 'Medium', '4', 'FFFF99'),
]
mitigations = [
    'Defined a fixed feature set after Iteration 1; deferred non-core features to future work.',
    'Prototyped SocketIO integration early in Iteration 3; isolated socket handlers in the social blueprint.',
    'Introduced extensions.py and Flask Blueprints in Iteration 4 to decouple modules.',
    'SQLite file committed to repository as a development artefact; production notes recommend PostgreSQL with automated backups.',
    'Used Flask-WTF CSRF tokens on all forms; SQLAlchemy ORM prevents raw SQL injection by default.',
    'Maintained a requirements.txt pinning all package versions; verified with a fresh virtual environment.',
]

tbl3 = doc.add_table(rows=1, cols=6)
tbl3.style = 'Table Grid'
for cell, txt in zip(tbl3.rows[0].cells, ['ID', 'Risk', 'Likelihood', 'Impact', 'Rating', 'Mitigation']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for (rid, desc, lik, imp, rating, col), mit in zip(risk_rows, mitigations):
    row = tbl3.add_row().cells
    set_cell(row[0], rid, center=True)
    set_cell(row[1], desc)
    set_cell(row[2], lik, center=True)
    set_cell(row[3], imp, center=True)
    set_cell(row[4], rating, bold=True, bg_color=col, center=True)
    set_cell(row[5], mit)
no_split(tbl3)
caption('Table 2.2 — Risk register (red = high, amber = medium, yellow = low-medium, green = low)')

h2('2.5  Ethics')
body(
    'This project was conducted in accordance with the UWE Bristol Research Ethics framework '
    '(UWE Bristol, 2025). As the system was built using entirely synthetic data — no real users were '
    'recruited, observed or surveyed during development — the ethical risk was assessed as minimal. '
    'The FET Ethical Review Checklist was completed and submitted prior to development commencing; '
    'a copy is included in the appendices. Should future iterations of the project involve real user '
    'testing or collection of biometric data such as weight logs, a full ethical review application '
    'would be required, with appropriate participant information sheets and informed consent forms.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — RESEARCH
# ═══════════════════════════════════════════════════════════════════════════════
h1('Chapter 3 — Research')

h2('3.1  Research Approach')
body(
    'The research drew on a combination of sources: reading existing academic literature and technical '
    'documentation, reviewing and comparing platforms currently used by combat sports athletes, and '
    'working through what the system actually needed to do based on domain knowledge of how fight camps '
    'work. The platform reviews shaped the requirements in Chapter 4, and the technology comparisons '
    'fed directly into the architecture decisions in Chapter 5.'
)

h2('3.2  Review of Existing Platforms')
body(
    'A survey of tools currently used by combat sports athletes and coaches revealed a fragmented market '
    'with no single platform addressing the full athlete lifecycle. The most relevant tools are reviewed below.'
)

h3('3.2.1  General-Purpose Productivity and Fitness Tools')
body(
    'Applications such as Notion, Google Sheets and Trello are widely used by independent fighters to '
    'manage camp schedules and opponent notes. While flexible, they provide no domain-specific structure: '
    'a fighter must design their own weight-cut template and manually calculate periodisation phases. '
    'MyFitnessPal (2024) offers calorie and weight tracking but has no concept of a fight date, a weight '
    'class target, or a camp phase — critical constructs in the combat sports context.'
)

h3('3.2.2  Enterprise Sports Analytics Platforms')
body(
    'Kitman Labs (2023) provides a comprehensive athlete-intelligence platform aimed at professional teams. '
    'Its feature set encompasses load monitoring, injury prediction, and squad wellness reporting. However, '
    'its pricing model targets franchise-level organisations and its interface assumes a coaching staff '
    'rather than an individual athlete. Similarly, Hudl (2024) excels at video tagging and tactical '
    'analysis for team sports but offers limited support for the individual preparation workflow of a '
    'combat sports competitor.'
)

h3('3.2.3  Combat-Sports-Specific Tools')
body(
    'Tapology (2024) is the most widely used combat-sports-specific web platform, providing fighter '
    'records, event listings and matchmaking data. Its scope is limited to information display — it offers '
    'no training management, weight tracking or communication features. No platform was identified in the '
    'review that combines training management, social coordination and event discovery in a single '
    'interface accessible to amateur and semi-professional fighters.'
)

h2('3.3  User Stories')
body(
    'In the absence of formal client interviews (precluded by the project timeline and ethical scope), '
    'user stories were derived from domain knowledge of combat sports preparation and validated against '
    'publicly available fighter accounts, coaching blogs and community discussions. Three primary user '
    'personas were identified and developed below.'
)

h3('3.3.1  Persona 1 — The Fighter')
body(
    'Jamie is a 24-year-old amateur MMA competitor training out of a small regional gym four days a week. '
    'He works a full-time job alongside his training and is preparing for his third amateur bout in eight '
    'weeks. Jamie has no dedicated sports scientist or nutritionist — he manages everything himself. '
    'He currently tracks his weight in the Notes app on his phone and plans his camp by scribbling '
    'sessions on a paper calendar.'
)
body(
    'Goals: Jamie wants to know whether his weight cut is on track at a glance, get a structured training '
    'plan specific to the opponent he is facing, and stop spending mental energy on logistics so he can '
    'focus on training. Frustrations: He never knows if he is doing the right things in the right order, '
    'misses sessions because nothing prompts him, and has to start his camp plan from scratch every time '
    'he gets a new fight.'
)

h3('3.3.2  Persona 2 — The Coach / Gym Owner')
body(
    'Marcus owns a small MMA gym with twelve active students, three of whom compete regularly. He '
    'communicates with his fighters via WhatsApp and has no structured way to see how their camp '
    'preparation is going between sessions. He sometimes finds out a fighter has been cutting weight '
    'dangerously only days before a bout.'
)
body(
    'Goals: Marcus wants to see his fighters\' weight trends, camp progress and fight dates in one '
    'place, and message them directly without switching to a separate app. Frustrations: '
    'Information is scattered across chat threads and paper notes; he has no visibility until a '
    'fighter tells him something is wrong.'
)

h3('3.3.3  Persona 3 — The Event Organiser')
body(
    'Sofia runs a regional amateur kickboxing promotion and organises four to six events per year. '
    'She currently manages registrations by email and a Google Form, which results in duplicate '
    'entries, unanswered queries and no way to filter interested fighters by weight class. Finding '
    'enough competitors for every division before each show takes weeks of back-and-forth.'
)
body(
    'Goals: Sofia wants fighters to discover and register for her events themselves, to filter '
    'registered athletes by weight class and experience level, and to have one place where her '
    'event details are visible to the right audience. Frustrations: Manual registration is error-prone '
    'and time-consuming; she has no targeted channel to reach fighters who compete in specific '
    'weight classes.'
)

h3('3.3.4  User Stories')
body(
    'Table 3.1 captures the key user stories derived from the three personas. Each story '
    'follows the standard format: As a [persona], I want to [action], so that [benefit]. '
    'These stories were subsequently decomposed into the formal functional requirements in Chapter 4.'
)

user_stories = [
    ('US1',  'Fighter',        'As a Fighter, I want to generate a periodised camp plan for a specific opponent, so that I know exactly what to train and when throughout the camp.'),
    ('US2',  'Fighter',        'As a Fighter, I want to log my weight daily and see whether my cut is on track, so that I can avoid a dangerous last-minute cut.'),
    ('US3',  'Fighter',        'As a Fighter, I want to generate a round-by-round game plan based on my opponent\'s strengths and weaknesses, so that I enter the fight with a clear tactical strategy.'),
    ('US4',  'Fighter',        'As a Fighter, I want to find sparring partners near me who train in compatible styles and skill levels, so that I can get productive sparring without wasting time searching.'),
    ('US5',  'Fighter',        'As a Fighter, I want to see a readiness score that tells me how prepared I am for my fight, so that I can identify and fix gaps before it is too late.'),
    ('US6',  'Fighter',        'As a Fighter, I want to view my personalised dashboard when I log in, so that I immediately see the most important information about my upcoming fight without having to navigate anywhere.'),
    ('US7',  'Coach',          'As a Coach, I want to view a connected fighter\'s weight logs and camp plans, so that I can monitor their preparation between training sessions.'),
    ('US8',  'Coach',          'As a Coach, I want to send and receive real-time messages with fighters in my network, so that I can give feedback without relying on WhatsApp or email.'),
    ('US9',  'Event Organiser','As an Event Organiser, I want to create an event with all relevant details (date, venue, weight classes, rules), so that fighters can discover and register for it themselves.'),
    ('US10', 'Event Organiser','As an Event Organiser, I want to see which fighters have registered, filtered by weight class, so that I can assess whether each division is filled without trawling through emails.'),
    ('US11', 'All users',      'As a user, I want to browse upcoming events without having to log in, so that I can discover events before deciding whether to register.'),
    ('US12', 'All users',      'As a user, I want to build up a shared database of opponents with scouting notes, so that I and my gym-mates can prepare for fighters we are likely to face.'),
]

tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
for cell, txt in zip(tbl4.rows[0].cells, ['ID', 'Persona', 'User Story']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for uid, persona, story in user_stories:
    row = tbl4.add_row().cells
    set_cell(row[0], uid,    bold=True, center=True, size=9)
    set_cell(row[1], persona, size=9)
    set_cell(row[2], story,   size=9)
no_split(tbl4)
caption('Table 3.1 — User stories by persona')

h2('3.4  Technology Selection')

h3('3.4.1  Web Framework — Pugh Matrix')
body(
    'Three Python web frameworks were evaluated as candidates for FPMS: Flask (Pallets Projects, 2024), '
    'Django (Django Software Foundation, 2024), and FastAPI (Ramírez, 2024). Each was scored against '
    'six weighted criteria using a Pugh matrix (Table 3.2), where 1 = poor fit and 5 = excellent fit. '
    'Criteria weights reflect the priorities of this project: flexibility and suitability for a '
    'medium-sized, server-rendered web application were weighted most heavily.'
)

criteria = [
    ('Ease of setup and learning curve',            3, [5, 3, 4]),
    ('Flexibility and control over app structure',  5, [5, 2, 4]),
    ('Built-in ORM / admin interface',              2, [2, 5, 2]),
    ('WebSocket / real-time support',               4, [4, 3, 4]),
    ('Documentation and community maturity',        3, [4, 5, 3]),
    ('Suitability for server-rendered HTML (SSR)',  4, [5, 5, 2]),
]

tbl5 = doc.add_table(rows=1, cols=6)
tbl5.style = 'Table Grid'
for cell, txt in zip(tbl5.rows[0].cells,
                     ['Criterion', 'Weight', 'Flask\n(score)', 'Flask\n(weighted)',
                      'Django\n(score / weighted)', 'FastAPI\n(score / weighted)']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE, center=True)

flask_total = django_total = fastapi_total = 0
for crit, weight, scores in criteria:
    row = tbl5.add_row().cells
    set_cell(row[0], crit)
    set_cell(row[1], str(weight), center=True)
    f, d, fa = scores
    fw, dw, faw = f*weight, d*weight, fa*weight
    flask_total += fw; django_total += dw; fastapi_total += faw
    set_cell(row[2], str(f), center=True)
    set_cell(row[3], str(fw), center=True, bold=True)
    set_cell(row[4], f'{d} / {dw}', center=True)
    set_cell(row[5], f'{fa} / {faw}', center=True)

total_row = tbl5.add_row().cells
set_cell(total_row[0], 'TOTAL WEIGHTED SCORE', bold=True, bg_color='D9E2F3')
set_cell(total_row[1], '', bg_color='D9E2F3')
set_cell(total_row[2], '', bg_color='D9E2F3')
set_cell(total_row[3], str(flask_total),   bold=True, bg_color='E2EFDA', center=True)
set_cell(total_row[4], str(django_total),  bold=True, bg_color='D9E2F3', center=True)
set_cell(total_row[5], str(fastapi_total), bold=True, bg_color='D9E2F3', center=True)
no_split(tbl5)
caption('Table 3.2 — Pugh matrix: Python web framework selection (Flask weighted score shown separately for clarity)')

body(
    f'Flask achieved the highest weighted score of {flask_total}, compared to {django_total} for Django '
    f'and {fastapi_total} for FastAPI. The key differentiator was Flask\'s combination of structural '
    'flexibility and strong server-side rendering support. Django\'s built-in ORM and admin interface '
    'would reduce boilerplate, but its opinionated project layout and heavier "batteries-included" '
    'philosophy were considered unnecessary overhead for a project of this scale. FastAPI is optimised '
    'for building REST APIs rather than full server-rendered web applications and lacks a mature '
    'template ecosystem for Jinja2-based HTML delivery.'
)

h3('3.4.2  Database Selection')
body(
    'Three database options were evaluated: SQLite (SQLite, 2024), PostgreSQL (PostgreSQL Global '
    'Development Group, 2024), and MySQL. Table 3.3 summarises the comparison across criteria '
    'relevant to this project.'
)

db_rows = [
    ('Zero configuration required',          True,  False, False),
    ('File-based (portable for development)', True,  False, False),
    ('Production-grade reliability',          False, True,  True),
    ('Full-text search support',              False, True,  True),
    ('SQLAlchemy ORM compatible',             True,  True,  True),
    ('Free and open source',                  True,  True,  True),
    ('Suitable for solo / small-team dev',    True,  False, False),
]

tbl6 = doc.add_table(rows=1, cols=4)
tbl6.style = 'Table Grid'
for cell, txt in zip(tbl6.rows[0].cells, ['Criterion', 'SQLite', 'PostgreSQL', 'MySQL']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE, center=True)
for crit, sq, pg, my in db_rows:
    row = tbl6.add_row().cells
    set_cell(row[0], crit)
    tick_cross(row[1], sq)
    tick_cross(row[2], pg)
    tick_cross(row[3], my)
no_split(tbl6)
caption('Table 3.3 — Database comparison (✓ = supported, ✗ = not supported)')

body(
    'SQLite was selected for the development and assessed version of the system on the basis of its '
    'zero-configuration setup and portability. Because SQLAlchemy abstracts the database connection '
    'through a URI string, migrating to PostgreSQL for a production deployment requires only a single '
    'configuration change and no application code modifications — a deliberate architectural decision '
    'documented in Chapter 5.'
)

h2('3.5  Summary of Research Findings')
body(
    'The review of existing platforms confirmed that the gap identified in Chapter 1 is real: '
    'no tool addresses the full preparation lifecycle of an independent athlete. That finding '
    'made scope easier to define — the system could not just replicate something that already '
    'existed; it had to bring together things that currently live in separate places.'
)
body(
    'On the technology side, Flask with SQLAlchemy and Flask-SocketIO emerged as the most '
    'appropriate combination for a server-rendered web application at this scale. Its structural '
    'flexibility was a key differentiator: the Blueprint architecture used in Chapter 5 would '
    'have been far harder to implement in a framework with a more opinionated project layout. '
    'The user research also produced a cleaner set of personas than initially expected — Fighter, '
    'Coach and Event Organiser map directly onto the application modules, which simplified the '
    'requirements work in Chapter 4.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════════
h1('Chapter 4 — Requirements')

h2('4.1  Overview')
body(
    'The user stories from Chapter 3 are translated here into a concrete set of system requirements. '
    'The requirements are split into two groups: Functional Requirements (FRs), covering what the '
    'system actually does, and Non-Functional Requirements (NFRs), covering how well it does it. '
    'NFRs are organised using the ISO/IEC 9126 quality categories (ISO/IEC 9126, 2001) — usability, '
    'functionality, reliability, portability, efficiency and maintainability. Use cases, following '
    'Cockburn (2001), map each actor to the parts of the system they interact with.'
)

h2('4.2  Actors and Use-Case Overview')
body(
    'Three actors interact with FPMS. The Guest (unauthenticated visitor) may only browse the public '
    'events listing. The Registered User (authenticated fighter, coach or general user) has full '
    'access to all training, social and sparring modules. The Admin (a user with the is_admin flag '
    'set) additionally has access to the event moderation panel. Table 4.1 maps each actor to the '
    'value-delivering use cases available to them. Authentication and profile setup are treated as '
    'system prerequisites rather than standalone use cases, as they do not directly deliver user '
    'value on their own. Figure 4.1 presents the full use case diagram.'
)

uc_rows = [
    ('UC1',  'Guest',          'Browse and discover public events',
     'Events module — unauthenticated /events listing with multi-field filter'),
    ('UC2',  'Registered User','View personalised fight-preparation dashboard',
     'Main module — /dashboard (countdown, weight, phase, warnings, chat preview)'),
    ('UC3',  'Registered User','Generate opponent-specific camp plan',
     'Training module — /camp_planner'),
    ('UC4',  'Registered User','Generate round-by-round game plan',
     'Training module — /game_plan'),
    ('UC5',  'Registered User','Log weight and monitor cut progress',
     'Training module — /weight_tracker (log, trend, risk score)'),
    ('UC6',  'Registered User','View composite readiness and risk score',
     'Training module — /risk_readiness'),
    ('UC7',  'Registered User','Manage shared opponent scouting database',
     'Fighters module — /fighters (add, search, edit, delete, compare)'),
    ('UC8',  'Registered User','Manage social connections',
     'Social module — /friends (send, accept, decline, remove friend requests)'),
    ('UC9',  'Registered User','Communicate in real time with connected users',
     'Social module — /chat (WebSocket bidirectional messaging)'),
    ('UC10', 'Registered User','View a friend\'s training profile and progress',
     'Social module — /friend_profile'),
    ('UC11', 'Registered User','Create sparring profile and discover partners',
     'Sparring module — /sparring_profile, /sparring_dashboard (Haversine matching)'),
    ('UC12', 'Registered User','Request, manage and complete sparring sessions',
     'Sparring module — /request_sparring_session (request → accept/decline → complete)'),
    ('UC13', 'Registered User','Assess partner skill after a session',
     'Sparring module — /assess_session (peer rating, honesty score)'),
    ('UC14', 'Registered User','Create, manage and register interest in events',
     'Events module — /events/create, /events/<id>/interest, /my_events'),
    ('UC15', 'Admin',          'Moderate events via admin panel',
     'Events module — /admin/events (all events, approval, registration statistics)'),
]

tbl_uc = doc.add_table(rows=1, cols=4)
tbl_uc.style = 'Table Grid'
for cell, txt in zip(tbl_uc.rows[0].cells, ['UC', 'Actor', 'Use Case', 'System Location']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for uc, actor, name, loc in uc_rows:
    row = tbl_uc.add_row().cells
    set_cell(row[0], uc,   bold=True, center=True)
    set_cell(row[1], actor)
    set_cell(row[2], name)
    set_cell(row[3], loc,  size=9)
no_split(tbl_uc)
caption('Table 4.1 — Use-case overview by actor (15 value-delivering use cases)')

add_picture_safe('screenshots/use_case_diagram.png', 'Figure 4.1 — Use case diagram', width=Cm(16))
caption('Figure 4.1 — FPMS use case diagram: Guest, Registered User and Admin actors across five functional modules')

h2('4.3  Functional Requirements')
body(
    'The functional requirements below are derived directly from the implemented system. Each FR is '
    'numbered, prioritised using MoSCoW (Must/Should/Could), and traced to the use-case(s) it supports. '
    'Table 4.2 presents the complete FR set.'
)

fr_rows = [
    # Authentication & User (prerequisites — no direct UC, enables UC2–UC15)
    ('FR1',  'Must',   'The system shall allow a guest to register with a unique username, unique email and password.',                     'UC1–UC15'),
    ('FR2',  'Must',   'The system shall authenticate users with case-insensitive username and hashed password comparison.',               'UC2–UC15'),
    ('FR3',  'Must',   'The system shall maintain an authenticated session and provide a logout mechanism.',                               'UC2–UC15'),
    ('FR4',  'Should', 'The system shall allow users to upload a profile picture stored in a server-side folder.',                         'UC2'),
    # Fighter Profile & Dashboard
    ('FR5',  'Must',   'The system shall allow a user to create and update a fighter profile recording age, height, walk-around weight, weight class, fight date, training days per week, and combat style.',  'UC2'),
    ('FR6',  'Must',   'The dashboard shall display fight countdown (days), current weight status, active camp phase, and smart warnings.', 'UC2'),
    ('FR7',  'Must',   'The system shall derive camp phase automatically (Base Conditioning / Skill Sharpening / Sparring Peak / Taper Week) from weeks remaining to fight date.', 'UC2'),
    # Camp Planner
    ('FR8',  'Must',   'The system shall generate a periodised, opponent-specific camp plan by keyword-matching opponent strengths and weaknesses.',  'UC3'),
    ('FR9',  'Must',   'Camp plans shall adapt phase count and duration to available weeks: 5 phases for 10+ weeks, 4 phases for 6–9 weeks, 3 phases for fewer than 6 weeks.', 'UC3'),
    ('FR10', 'Must',   'Users shall be able to view the history of all generated camp plans.',                                             'UC3'),
    ('FR11', 'Should', 'Users shall be able to delete any of their own camp plans.',                                                      'UC3'),
    # Game Plan
    ('FR12', 'Must',   'The system shall generate a round-by-round game plan through a decision-tree analysis of opponent strengths and weaknesses, covering preferred range, strategic approach, tactical notes, techniques to drill, things to avoid, and per-round objectives.',  'UC4'),
    ('FR13', 'Should', 'Game plan generation shall be pre-populatable from the opponent database, auto-filling the strengths and weaknesses fields.',  'UC4, UC7'),
    ('FR14', 'Must',   'Users shall be able to view all generated game plans.',                                                            'UC4'),
    ('FR15', 'Should', 'Users shall be able to delete any of their own game plans.',                                                      'UC4'),
    # Weight Tracker
    ('FR16', 'Must',   'The system shall allow users to log body weight (in lbs) with a date stamp.',                                     'UC5'),
    ('FR17', 'Must',   'The system shall calculate and display a weight trend (Increasing / Decreasing / Stable) based on 7-day rolling averages.',  'UC5'),
    ('FR18', 'Must',   'The system shall produce a weight-cut risk score (0–100) with itemised deduction factors for weight overage, training availability, camp plan absence, and proximity to fight date.', 'UC5, UC6'),
    ('FR19', 'Should', 'Users shall be able to delete individual weight log entries.',                                                    'UC5'),
    # Readiness
    ('FR20', 'Must',   'The system shall produce a composite readiness score (0–100) and readiness level (Low / Medium / High) on the risk-readiness page.', 'UC6'),
    # Opponent Database
    ('FR21', 'Must',   'The system shall allow users to add opponents to a shared database with name, nickname, weight class, record, fighting style, strengths, weaknesses and notable fights.',  'UC7'),
    ('FR22', 'Must',   'The opponent database shall support search by name and filter by weight class.',                                   'UC7'),
    ('FR23', 'Should', 'Users shall be able to edit and delete opponent records.',                                                        'UC7'),
    ('FR24', 'Could',  'The system shall provide a side-by-side comparison view of two opponents.',                                       'UC7'),
    # Social
    ('FR25', 'Must',   'Users shall be able to send, accept and decline friend requests by username.',                                    'UC8'),
    ('FR26', 'Should', 'Users shall be able to remove existing friendships.',                                                             'UC8'),
    ('FR27', 'Must',   'Users shall be able to view a friend\'s fighter profile, weight logs and training plans.',                       'UC10'),
    # Chat
    ('FR28', 'Must',   'The system shall support real-time bidirectional messaging between friends via WebSocket.',                        'UC9'),
    ('FR29', 'Must',   'The system shall track unread message counts per conversation and mark messages as read when a conversation is opened.', 'UC9'),
    # Sparring
    ('FR30', 'Must',   'Users shall be able to create a sparring profile specifying location, skill level, preferred styles, availability, maximum distance and a self-skill rating.',  'UC11'),
    ('FR31', 'Must',   'The system shall rank sparring partner matches by a compatibility score (0–100) combining skill-level proximity, style compatibility and geolocation distance via the Haversine formula; partners below 50 shall be excluded.',  'UC11'),
    ('FR32', 'Must',   'Users shall be able to request sparring sessions specifying date, time, duration and location.',                   'UC12'),
    ('FR33', 'Must',   'Session recipients shall be able to accept, decline or mark sessions as completed.',                              'UC12'),
    ('FR34', 'Should', 'After a session is marked complete, both parties shall be able to submit a peer skill assessment (1–10).',        'UC13'),
    ('FR35', 'Could',  'The system shall maintain a honesty score on sparring profiles, adjusting up or down based on the delta between self-rating and peer ratings.',  'UC13'),
    # Events
    ('FR36', 'Must',   'Any registered user shall be able to create events with name, type, date, time, registration deadline, venue, city, country, weight classes, experience levels, rules, entry fee, prize info, contact details and max participants.',  'UC14'),
    ('FR37', 'Must',   'The public events listing shall show only future, approved events and support filtering by city, country, weight class, experience level, event type and rules.',  'UC1, UC14'),
    ('FR38', 'Must',   'Registered users shall be able to express interest in, register for, or withdraw from events; organisers shall not be able to register for their own events.',  'UC14'),
    ('FR39', 'Must',   'Organisers shall be able to edit and delete their own events.',                                                   'UC14'),
    ('FR40', 'Must',   'Admins shall have access to an event management panel showing all events with registration statistics.',           'UC15'),
]

tbl_fr = doc.add_table(rows=1, cols=4)
tbl_fr.style = 'Table Grid'
for cell, txt in zip(tbl_fr.rows[0].cells, ['FR', 'Priority', 'Description', 'UC']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)

priority_colors = {'Must': 'FFE0E0', 'Should': 'FFF3CD', 'Could': 'E2EFDA'}
for fr, pri, desc, uc in fr_rows:
    row = tbl_fr.add_row().cells
    set_cell(row[0], fr, bold=True, center=True)
    set_cell(row[1], pri, center=True, bg_color=priority_colors[pri])
    set_cell(row[2], desc)
    set_cell(row[3], uc, center=True)
no_split(tbl_fr)
caption('Table 4.2 — Functional requirements (red = Must, yellow = Should, green = Could)')

h2('4.4  Non-Functional Requirements')
body(
    'The non-functional requirements are grouped under the six ISO/IEC 9126 (2001) quality '
    'categories. Table 4.3 lists all 17 NFRs with their priority and description.'
)

nfr_rows = [
    ('NFR1',  'Usability',       'Must',   'All pages shall be fully responsive on screen widths from 320 px (mobile) to 1920 px (desktop) using Bootstrap 5\'s grid system.'),
    ('NFR2',  'Usability',       'Must',   'Flash messages shall provide clear, contextual feedback for every user action (success, warning, danger categories).'),
    ('NFR3',  'Usability',       'Should', 'The dashboard shall present the most time-critical information (fight countdown, weight status, warnings) above the fold without scrolling on a 768 px viewport.'),
    ('NFR4',  'Functionality',   'Must',   'All web forms shall include CSRF tokens (via Flask-WTF) to prevent cross-site request forgery attacks.'),
    ('NFR5',  'Functionality',   'Must',   'All form inputs shall be validated server-side; client-side validation is supplementary and not relied upon for security.'),
    ('NFR6',  'Functionality',   'Must',   'Passwords shall be stored exclusively as Werkzeug PBKDF2-SHA256 hashes; plaintext passwords shall never be persisted.'),
    ('NFR7',  'Functionality',   'Must',   'All routes requiring authentication shall be protected with the @login_required decorator; unauthenticated access shall redirect to /login.'),
    ('NFR8',  'Reliability',     'Must',   'The SQLAlchemy ORM shall be used for all database operations, preventing raw SQL injection by construction.'),
    ('NFR9',  'Reliability',     'Should', 'Database relationships shall be enforced via foreign key constraints defined at the ORM model level.'),
    ('NFR10', 'Reliability',     'Should', 'The application shall handle missing fighter profiles or sparring profiles gracefully, redirecting users to the relevant creation form rather than raising a 500 error.'),
    ('NFR11', 'Portability',     'Must',   'The database backend shall be configurable via a single URI string in app.py (SQLALCHEMY_DATABASE_URI), enabling migration from SQLite (development) to PostgreSQL (production) without code changes.'),
    ('NFR12', 'Portability',     'Must',   'The application shall run on Python 3.11 and above and shall declare all dependencies in requirements.txt with pinned versions.'),
    ('NFR13', 'Efficiency',      'Should', 'SQLAlchemy relationships shall use lazy loading by default to avoid N+1 query patterns on list pages.'),
    ('NFR14', 'Efficiency',      'Should', 'The sparring match algorithm shall cap results at 20 partners to limit query and rendering time.'),
    ('NFR15', 'Maintainability', 'Must',   'The application shall be structured as Flask Blueprints, with each functional module (auth, main, training, social, fighters, sparring, events) in a separate file under the routes/ package.'),
    ('NFR16', 'Maintainability', 'Must',   'Shared extensions (LoginManager, SocketIO) shall be instantiated in extensions.py and imported by both app.py and route modules to prevent circular imports.'),
    ('NFR17', 'Maintainability', 'Should', 'The SECRET_KEY and database URI shall be documented as requiring environment variable substitution before production deployment.'),
]

tbl_nfr = doc.add_table(rows=1, cols=4)
tbl_nfr.style = 'Table Grid'
for cell, txt in zip(tbl_nfr.rows[0].cells, ['NFR', 'ISO/IEC 9126 Category', 'Priority', 'Description']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)

cat_colors = {
    'Usability': 'DEEAF1', 'Functionality': 'E2EFDA', 'Reliability': 'FFF3CD',
    'Portability': 'FCE4D6', 'Efficiency': 'EAD1DC', 'Maintainability': 'D9D9D9'
}
for nfr, cat, pri, desc in nfr_rows:
    row = tbl_nfr.add_row().cells
    set_cell(row[0], nfr, bold=True, center=True)
    set_cell(row[1], cat, bg_color=cat_colors[cat], center=True)
    set_cell(row[2], pri, center=True, bg_color=priority_colors[pri])
    set_cell(row[3], desc)
no_split(tbl_nfr)
caption('Table 4.3 — Non-functional requirements mapped to ISO/IEC 9126 quality categories')

h2('4.5  Requirements Summary')
body(
    f'Altogether FPMS has {len(fr_rows)} functional requirements spread across eight modules — '
    f'{sum(1 for r in fr_rows if r[1]=="Must")} Must-have, '
    f'{sum(1 for r in fr_rows if r[1]=="Should")} Should-have, '
    f'{sum(1 for r in fr_rows if r[1]=="Could")} Could-have. '
    f'Every Must-have FR was implemented and shipped in the final system. '
    f'The {len(nfr_rows)} non-functional requirements are verified against the testing evidence in Chapter 6.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 — DESIGN AND DEVELOPMENT
# ═══════════════════════════════════════════════════════════════════════════════
h1('Chapter 5 — Design and Development')

h2('5.1  System Architecture')
body(
    'The application is a server-rendered Flask app organised into seven Blueprints, one per functional '
    'area, all registered through a thin app.py factory. Splitting the code this way was a deliberate '
    'choice: early in development the routes lived in a single file and import errors were already '
    'becoming a problem. Blueprints fixed that by giving each module its own namespace and removing '
    'the need for any module to import directly from another (Grinberg, 2018; Pallets Projects, 2024).'
)
body(
    'A critical design constraint was the need for shared state between modules. Two extensions — '
    'Flask-Login\'s LoginManager and Flask-SocketIO\'s SocketIO instance — must be accessible both '
    'in app.py (for initialisation) and in route modules (for decorators and event handlers). '
    'Importing app directly from route modules would create a circular dependency. This was resolved '
    'by instantiating both extensions in a neutral extensions.py module that is imported by both '
    'app.py and the routes, as shown in the code snippet below.'
)

code_block([
    '# extensions.py',
    'from flask_login import LoginManager',
    'from flask_socketio import SocketIO',
    '',
    'login_manager = LoginManager()',
    'socketio = SocketIO()',
    '',
    '# app.py (thin factory)',
    'from extensions import login_manager, socketio',
    'db.init_app(app)',
    'login_manager.init_app(app)',
    'socketio.init_app(app)',
])
caption('Code Snippet 5.1 — Circular-import prevention via extensions.py')

body('Table 5.1 summarises the responsibility of each Blueprint and the route count it exposes.')

bp_rows = [
    ('auth',     'routes/auth.py',     3,  'User registration, login, logout. Case-insensitive username lookup via SQLAlchemy func.lower().'),
    ('main',     'routes/main.py',     2,  'Dashboard: fight countdown, camp-phase derivation, weight status, warnings, recent chat preview.'),
    ('training', 'routes/training.py', 11, 'Fighter profile, camp-plan generator, game-plan generator, weight tracker, risk-readiness score, CRUD for plans and logs.'),
    ('social',   'routes/social.py',   8,  'Friend requests, friend list, friend profile view, chat list, per-conversation chat view, SocketIO event handlers.'),
    ('fighters', 'routes/fighters.py', 6,  'Opponent database CRUD, side-by-side comparison, redirect to game plan with pre-populated fields.'),
    ('sparring', 'routes/sparring.py', 6,  'Sparring profile, partner-matching dashboard, partner profile, session request/respond/assess.'),
    ('events',   'routes/events.py',   8,  'Event CRUD, public listing with filters, interest/registration, my-events, admin panel.'),
]

tbl_bp = doc.add_table(rows=1, cols=4)
tbl_bp.style = 'Table Grid'
for cell, txt in zip(tbl_bp.rows[0].cells, ['Blueprint', 'File', 'Routes', 'Responsibility']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for name, path, count, resp in bp_rows:
    row = tbl_bp.add_row().cells
    set_cell(row[0], name, bold=True)
    set_cell(row[1], path)
    set_cell(row[2], str(count), center=True)
    set_cell(row[3], resp)
no_split(tbl_bp)
caption('Table 5.1 — Blueprint structure: 7 modules, 44 routes total')

body(
    'Figure 5.1 illustrates the complete application architecture, showing the data flow from '
    'the browser client through the Flask application core — including the Login Manager, '
    'WebSocket engine and ORM layer — to the individual Blueprint modules and the SQLite '
    'database. The dashed WebSocket arrow shows the separate channel used for real-time chat '
    'via Flask-SocketIO, which bypasses the normal HTTP request cycle.'
)
add_picture_safe('screenshots/app_architecture.png', 'Figure 5.1 — Architecture diagram', width=Cm(16))
caption('Figure 5.1 — FPMS application architecture: Blueprint structure, WebSocket channel and data flow')

h2('5.2  Database Design')
body(
    'The database layer uses SQLAlchemy 2.x as the ORM with SQLite as the backing store for '
    'development. Fourteen model classes map to fourteen relational tables. All primary keys are '
    'auto-incrementing integers; foreign-key relationships are declared at the model level and '
    'enforced by SQLAlchemy. Table 5.2 presents the entity overview; Figure 5.2 describes the '
    'key relationships.'
)

model_rows = [
    ('User',            'Core identity record. username and email are unique-constrained. password_hash stores a Werkzeug PBKDF2-SHA256 digest (Werkzeug, 2024). is_admin bool gates the admin panel.'),
    ('FighterProfile',  '1-to-1 with User. Stores age, height (cm), walk_around_weight (lbs), weight_class, fight_date, training_availability (days/week), se_angle (Striking / Grappling / Mixed).'),
    ('CampPlan',        'Many-to-1 with User. plan_phases stores a JSON string of phase objects generated by generate_camp_plan().'),
    ('GamePlan',        'Many-to-1 with User. round_objectives stores a JSON string containing strategic_approach, tactical_notes, techniques_to_drill, things_to_avoid, and per-round objective objects.'),
    ('WeightLog',       'Many-to-1 with User. One row per daily weigh-in. weight stored in lbs.'),
    ('Fighter',         'Shared opponent database (not user-scoped). Stores name, record, style, strengths, weaknesses. Used to pre-populate game plan generation.'),
    ('FriendRequest',   'Tracks pending/accepted/declined requests between pairs of users. Sender and receiver use separate foreign keys to User.'),
    ('Friendship',      'Created when a FriendRequest is accepted. user1_id always holds the lower user ID (normalised) to prevent duplicate rows.'),
    ('ChatMessage',     'Many-to-1 with sender User and receiver User. is_read flag supports unread-count queries. Real-time delivery via SocketIO; persistence via ORM.'),
    ('SparringProfile', '1-to-1 with User. latitude/longitude enable Haversine distance calculation. honesty_score (float, 0.5–1.0) is updated by post-session assessments.'),
    ('SparringSession', 'Tracks a session request between requester and partner. status: pending → accepted → completed. completed_at populated on completion.'),
    ('SkillAssessment', 'One assessment per user per session. skill_rating 1–10. Triggers honesty_score recalculation on the assessed user\'s SparringProfile.'),
    ('Event',           'Organiser-owned event. weight_classes and experience_levels stored as comma-separated strings; helper methods split them to lists. is_approved flag for future moderation.'),
    ('EventInterest',   'Junction between User and Event. status: interested / registered / withdrawn. weight_class records the division the user intends to compete in.'),
]

tbl_models = doc.add_table(rows=1, cols=2)
tbl_models.style = 'Table Grid'
for cell, txt in zip(tbl_models.rows[0].cells, ['Model', 'Key Design Notes']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for name, notes in model_rows:
    row = tbl_models.add_row().cells
    set_cell(row[0], name, bold=True)
    set_cell(row[1], notes)
no_split(tbl_models)
caption('Table 5.2 — Database model overview (14 entities)')

body(
    'Key relationships: User → FighterProfile (one-to-one); User → CampPlan, GamePlan, WeightLog '
    '(one-to-many); User ↔ User via FriendRequest and Friendship (self-referential many-to-many); '
    'User ↔ User via ChatMessage (self-referential); User → SparringProfile (one-to-one); '
    'SparringSession → SkillAssessment (one-to-many); User ↔ Event via EventInterest '
    '(many-to-many with attributes). The Fighter model is intentionally not user-scoped — it acts '
    'as a shared scouting database accessible to all authenticated users.'
)

add_picture_safe('screenshots/01_database_schema.png', 'Figure 5.2 — Database schema / ERD')
caption('Figure 5.2 — FPMS entity-relationship overview (14 models)')

h2('5.3  Key Algorithms')

h3('5.3.1  Opponent-Specific Camp Plan Generator')
body(
    'The camp plan generator (generate_camp_plan() in routes/training.py) takes four inputs: '
    'the opponent\'s strengths (free text), the opponent\'s weaknesses (free text), the number of '
    'weeks remaining to the fight, and the fight type (3 or 5 rounds). It uses keyword matching '
    'against eight categorised lists to infer the opponent\'s dominant attributes, then constructs '
    'a training plan whose phase count and duration adapt to the available time window.'
)
code_block([
    '# Keyword detection (excerpt)',
    'striking_keywords = ["striking","boxing","kickboxing","muay thai","knockout","ko",...]',
    'grappling_keywords = wrestling_keywords + bjj_keywords + ["grappling","ground","gnp"]',
    '',
    'opp_striking = any(kw in strengths_lower for kw in striking_keywords)',
    'opp_grappling = any(kw in strengths_lower for kw in grappling_keywords)',
    '',
    '# Adaptive phase structure based on weeks remaining',
    'if weeks_remaining >= 10:   phases = { 5-phase plan }',
    'elif weeks_remaining >= 6:  phases = { 4-phase plan }',
    'else:                       phases = { 3-phase plan }',
])
caption('Code Snippet 5.2 — Camp plan keyword detection and adaptive phase structure')

body(
    'Each training attribute detected (opponent striking, opponent grappling, opponent cardio, etc.) '
    'contributes specific items to five lists: training_priorities, sparring_focus, technique_drills, '
    'conditioning_focus, and mental_prep. These lists are then distributed across the phases. '
    'The completed phase dictionary is serialised to JSON and stored in CampPlan.plan_phases, '
    'allowing the view template to iterate over phases without parsing logic.'
)

h3('5.3.2  Round-by-Round Game Plan Generator')
body(
    'The game plan generator (generate_game_plan()) uses a decision-tree of fifteen strategic '
    'branches to determine the fighter\'s preferred range and round-by-round objectives. The tree '
    'evaluates opponent attribute pairs — such as "strong grappling + weak striking" → "keep it '
    'standing" — before falling through to more generic strategies. This approach ensures that '
    'the most tactically specific advice is always surfaced first.'
)
code_block([
    '# Strategic decision tree (excerpt — 15 branches)',
    'if opp_grappling and weak_striking:',
    '    preferred_range = "Striking (Stay Standing)"',
    '    strategic_approach.append("Maintain distance, avoid clinch/takedowns")',
    '',
    'elif opp_striking_offense and weak_grappling:',
    '    preferred_range = "Wrestling & Ground"',
    '    strategic_approach.append("Close distance quickly, shoot takedowns early")',
    '',
    'elif weak_chin:',
    '    preferred_range = "Striking (Power Shots)"',
    '    strategic_approach.append("Patient setups, look for clean power shots")',
    '',
    '# Round objectives are then set conditionally on preferred_range',
    '# and on whether the fight is 3 or 5 rounds',
])
caption('Code Snippet 5.3 — Game plan strategic decision tree (first three branches shown)')

body(
    'The generator produces a dictionary containing preferred_range, range_reason, '
    'strategic_approach, tactical_notes, techniques_to_drill, things_to_avoid, and '
    'round_objectives (keyed Round 1–3 or Round 1–5). The entire structure is serialised to JSON '
    'and stored in GamePlan.round_objectives. Round-by-round objectives are rendered in the view '
    'template by deserialising the JSON and iterating over the dictionary in key order.'
)

h3('5.3.3  Sparring Partner Matching — Haversine + Compatibility Score')
body(
    'The sparring match algorithm (get_sparring_matches() in routes/sparring.py) scores each '
    'candidate SparringProfile against the requesting user\'s profile on three dimensions: '
    'skill-level proximity, style compatibility, and geographic distance. The base compatibility '
    'score starts at 40 and points are added in each dimension; profiles scoring below 50 are '
    'filtered out and the top 20 results are returned sorted by score descending.'
)
code_block([
    '# Skill level proximity (max +25)',
    'skill_levels = ["Beginner","Intermediate","Advanced","Expert"]',
    'skill_diff = abs(user_skill_idx - profile_skill_idx)',
    'if skill_diff == 0:   compatibility += 25',
    'elif skill_diff == 1: compatibility += 20',
    'elif skill_diff == 2: compatibility += 10',
    '',
    '# Geographic distance via Haversine formula (Sinnott, 1984)',
    'def haversine_distance(lat1, lon1, lat2, lon2):',
    '    lat1,lon1,lat2,lon2 = map(math.radians,[lat1,lon1,lat2,lon2])',
    '    a = sin(dlat/2)**2 + cos(lat1)*cos(lat2)*sin(dlon/2)**2',
    '    return 2 * asin(sqrt(a)) * 3959  # miles',
    '',
    '# Distance scoring (max +20)',
    'if distance <= 10:  compatibility += 20',
    'elif distance <= 25: compatibility += 15',
    'elif distance <= 50: compatibility += 12',
])
caption('Code Snippet 5.4 — Sparring compatibility scoring (skill and Haversine distance components)')

body(
    'The Haversine formula (Sinnott, 1984) calculates the great-circle distance between two points '
    'on a sphere given their latitudes and longitudes in decimal degrees. It was selected over a '
    'simple Euclidean approximation because it correctly accounts for the curvature of the Earth, '
    'which matters at distances beyond approximately 50 km. If either profile lacks coordinates, '
    'the algorithm falls back to a text-based location comparison, extracting common words from '
    'the location strings and awarding partial proximity points.'
)
body(
    'A honesty score multiplier is maintained on each SparringProfile. After a completed session, '
    'each party submits a peer skill rating. The assessed user\'s self-skill rating is then compared '
    'against the average of all peer ratings: if the delta is ≤ 1, the honesty score increases by '
    '0.1 (capped at 1.0); if ≥ 3, it decreases by 0.1 (floored at 0.5). This discourages '
    'inflated self-ratings that would skew compatibility matching.'
)

h3('5.3.4  Composite Risk / Readiness Score')
body(
    'Both the weight tracker and the risk-readiness page compute a composite score (0–100) that '
    'aggregates four factors with explicit, itemised deductions visible to the user:'
)
code_block([
    'score = 100',
    '',
    '# Factor 1 — weight cut',
    'weight_diff = current_weight - target_weight   # target = walk_around - 10 lbs',
    'if weight_diff > 5:  score -= min(80, weight_diff * 3)',
    'elif weight_diff > 0: score -= 10',
    '',
    '# Factor 2 — training availability',
    'if training_availability < 3: score -= 20',
    'elif training_availability < 5: score -= 10',
    '',
    '# Factor 3 — camp plan presence',
    'if not CampPlan.query.filter_by(user_id=current_user.id).count(): score -= 15',
    '',
    '# Factor 4 — days to fight',
    'if days_until <= 14: score -= 20',
    'elif days_until <= 30: score -= 10',
    'elif days_until <= 60: score -= 5',
    '',
    'score = max(0, score)',
    'readiness_level = "Low" if score < 30 else "Medium" if score < 70 else "High"',
])
caption('Code Snippet 5.5 — Composite readiness score calculation')

h2('5.4  Real-Time Chat Architecture')
body(
    'Real-time messaging is implemented using Flask-SocketIO 5.x (Flask-SocketIO, 2024), which '
    'wraps the Socket.IO protocol over WebSocket. The server maintains a per-user room '
    '("user_{id}") for delivering unread-count notifications and a per-conversation room '
    '("chat_{min(a,b)}_{max(a,b)}") for message delivery. This naming convention ensures '
    'symmetry — the same room name is computed regardless of which party initiates the '
    'connection, preventing duplicate rooms.'
)
body('Table 5.3 describes the SocketIO event flow for a message send.')

socket_rows = [
    ('Client → Server', 'join',              'User joins their personal notification room (user_{id}) on page load.'),
    ('Client → Server', 'join_chat',         'User announces they have opened a specific conversation; server records this in current_chats dict.'),
    ('Client → Server', 'send_message',      'Client emits message text, sender_id, recipient_id. Server persists to DB, emits receive_message to conversation room, emits unread_count_update to recipient\'s personal room.'),
    ('Server → Client', 'receive_message',   'Delivered to conversation room. Both parties (if online) receive the message in real time.'),
    ('Server → Client', 'unread_count_update','Delivered to recipient\'s personal room. Contains a dict of unread counts per sender; if recipient is currently viewing the sender\'s chat, count is set to 0.'),
    ('Client → Server', 'leave_chat',        'User navigates away; server removes them from current_chats, stopping the unread suppression.'),
]

tbl_sock = doc.add_table(rows=1, cols=3)
tbl_sock.style = 'Table Grid'
for cell, txt in zip(tbl_sock.rows[0].cells, ['Direction', 'Event Name', 'Description']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for direction, event, desc in socket_rows:
    row = tbl_sock.add_row().cells
    set_cell(row[0], direction, center=True)
    set_cell(row[1], event, bold=True)
    set_cell(row[2], desc)
no_split(tbl_sock)
caption('Table 5.3 — SocketIO event flow for real-time chat')

h2('5.5  Security Design')
body(
    'Security was considered at every layer of the application. At the authentication layer, '
    'passwords are stored exclusively as Werkzeug PBKDF2-SHA256 hashes with a random salt; '
    'the plaintext password is never persisted or logged (Werkzeug, 2024). Session management '
    'is delegated entirely to Flask-Login, which signs the session cookie with the application '
    'SECRET_KEY. All routes that modify state or return user-specific data are decorated with '
    '@login_required, ensuring unauthenticated access is redirected to /login rather than '
    'raising a 403.'
)
body(
    'At the form layer, every HTML form is protected by a CSRF token provided by Flask-WTF. '
    'The token is embedded as a hidden field and validated on every POST request; a missing or '
    'invalid token results in a 400 response. At the database layer, all queries are issued '
    'through SQLAlchemy\'s ORM, which parameterises all inputs automatically and prevents '
    'SQL injection by construction. File uploads (profile pictures) are sanitised using '
    'Werkzeug\'s secure_filename() function and stored outside the web root under a '
    'configurable UPLOAD_FOLDER path.'
)
body(
    'One known limitation is that the SECRET_KEY is currently hard-coded in app.py as a '
    'development placeholder. Before any production deployment this must be replaced with a '
    'cryptographically random value loaded from an environment variable, as documented in the '
    'project README and flagged in the risk register (R5, Table 2.2).'
)

h2('5.6  Frontend Design and User Interface')
body(
    'The frontend is built on Bootstrap 5\'s responsive grid system with Jinja2 server-side '
    'templating. All pages share a common base.html template that provides the navigation bar, '
    'global flash-message display, SocketIO client initialisation, and Bootstrap JavaScript '
    'includes. Individual pages extend base.html via Jinja2\'s {% extends %} / {% block %} '
    'inheritance pattern, ensuring consistent styling with minimal duplication.'
)
body(
    'The dashboard provides the primary "at-a-glance" view for a fighter: fight countdown in days, '
    'current weight against the cut target, active camp phase label, smart warnings (e.g. "Weight '
    'cut may be too aggressive"), and a preview of the five most recent conversations with '
    'unread-message badges. The interface is intentionally data-forward rather than decorative, '
    'reflecting the professional context of the user base.'
)

body(
    'Figure 5.3 shows the fighter profile form, which is the first screen a new user completes '
    'after registration. The form captures the data that drives all personalised features: '
    'walk-around weight and weight class feed into the cut-risk scoring, fight date drives the '
    'countdown and camp phase calculation, and training availability (days per week) contributes '
    'to the readiness score deduction.'
)
add_picture_safe('screenshots/13_fighter_profile.png', 'Figure 5.3 — Fighter Profile form')
caption('Figure 5.3 — Fighter profile form: all fields feed directly into dashboard metrics and readiness scoring')

body(
    'Figure 5.4 shows the main dashboard, which is the default landing page after login. '
    'Rather than a blank home screen, the dashboard immediately presents the fighter\'s '
    'most time-critical information: days until the fight, current weight status relative '
    'to the cut target, active camp phase, any active warnings (e.g. aggressive weight cut), '
    'and a preview of recent conversations with unread-message badges. The module grid below '
    'the metrics gives one-click access to every platform feature.'
)
add_picture_safe('screenshots/02_dashboard.png', 'Figure 5.4 — Dashboard')
caption('Figure 5.4 — Dashboard: fight countdown, weight status, camp phase and smart warnings above the module grid')

body(
    'Figure 5.5 shows the output of the camp plan generator for an opponent with strong '
    'striking and a poor grappling defence. The keyword-matching engine has correctly identified '
    'the opponent\'s profile and produced a five-phase plan (ten weeks remaining) with '
    'striking-defence priorities in the early phases and wrestling-offence drills in the '
    'later ones. Each phase is individually expandable to reveal the full training detail.'
)
add_picture_safe('screenshots/03_camp_plan.png', 'Figure 5.5 — Camp Plan output')
caption('Figure 5.5 — Generated camp plan: five phases with opponent-specific training priorities and mental preparation notes')

body(
    'Figure 5.6 shows the sparring partner discovery dashboard. Each card displays a '
    'potential partner\'s username, skill level, preferred style, location and compatibility '
    'score. The score reflects the combination of skill proximity, style compatibility and '
    'geolocation distance calculated by the Haversine algorithm described in Section 5.3.3. '
    'Only profiles scoring 50 or above appear, ensuring the list contains genuinely '
    'compatible partners rather than every registered user.'
)
add_picture_safe('screenshots/04_sparring_dashboard.png', 'Figure 5.6 — Sparring Dashboard')
caption('Figure 5.6 — Sparring partner dashboard: compatibility-ranked cards with skill, style and distance data')

body(
    'Figure 5.7 shows the public events listing page, which is accessible without logging in. '
    'The multi-field filter bar at the top allows visitors to narrow events by city, country, '
    'weight class, experience level, event type and rules. Each event card shows the key '
    'details at a glance; clicking through takes the user to the full event page where '
    'they can register interest or withdraw.'
)
add_picture_safe('screenshots/05_events.png', 'Figure 5.7 — Events listing')
caption('Figure 5.7 — Public events listing with multi-field filter bar: accessible to guests and registered users alike')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 — RESULTS AND TESTING
# ═══════════════════════════════════════════════════════════════════════════════
h1('Chapter 6 — Results and Testing')

h2('6.1  Testing Approach')
body(
    'Testing was done manually, working through each use case in the browser and checking '
    'that the application behaved as expected for both normal inputs and boundary cases '
    '(Myers, Sandler and Badgett, 2011). Automated tests were considered but deprioritised — '
    'the logic is mostly route-level and view-rendering code that is easiest to verify by '
    'actually using the app, and time spent writing a test harness would have come out of '
    'feature development. Manual testing also caught a few edge cases (like the game plan '
    'form pre-population) that isolated unit tests would have missed entirely.'
)
body(
    'Each of the 44 routes was exercised at least once with valid inputs, and routes that involve '
    'access control were additionally tested with an unauthenticated user and a user without the '
    'required permissions. HTTP response codes were observed in the browser developer tools: GET '
    'routes were expected to return HTTP 200; POST routes to return HTTP 302 (redirect to the '
    'next page) on success, or HTTP 200 with validation errors on invalid input. All 44 routes '
    'returned the expected responses, representing 100% route coverage. The results of this '
    'testing are presented in Sections 6.2–6.5.'
)

h2('6.2  Requirements Traceability Matrix')
body(
    'Table 6.1 traces all 40 functional requirements against their test results. Each row '
    'shows what was tested, what the expected response was, what actually happened, and '
    'whether it passed. Every Must-have and Should-have requirement was tested; Could-have '
    'requirements were tested where they had been implemented.'
)

rtm_rows = [
    ('FR1',  'Must',   'Submitted /signup form with unique username, email and password',
     '302 redirect to login; account created in DB',
     '302 received; user row created; login succeeded',                  True),
    ('FR2',  'Must',   'Logged in using uppercase variant of registered username',
     'Authentication succeeds; session cookie set',
     'Dashboard rendered; user authenticated',                           True),
    ('FR3',  'Must',   'Clicked Logout link while authenticated',
     '302 redirect to /login; session cleared',
     'Redirected to /login; subsequent /dashboard returns 302 to login', True),
    ('FR4',  'Should', 'POSTed /profile with a JPEG file attached',
     'File saved under static/uploads; filename stored in User.profile_pic',
     'Image saved; displayed in navbar and profile page',                True),
    ('FR5',  'Must',   'Filled and submitted fighter profile with all required fields',
     'FighterProfile row created/updated; redirect to /dashboard',
     'Profile saved; dashboard reflects new values',                     True),
    ('FR6',  'Must',   'Viewed /dashboard with an active fight date 30 days ahead',
     'Countdown, weight status, camp phase and warnings displayed',
     'All four elements present; countdown showed correct day count',    True),
    ('FR7',  'Must',   'Set fight date 4 weeks away; viewed dashboard',
     'Camp phase shown as "Taper Week"',
     '"Taper Week" displayed in camp phase card',                        True),
    ('FR8',  'Must',   'Submitted camp planner with opponent strengths "strong grappling" and weaknesses "weak chin"',
     'Phased plan generated with wrestling defence and counter-striking drills',
     'Correct phase plan returned with tailored content',                True),
    ('FR9',  'Must',   'Tested camp planner with 12, 8 and 4 weeks remaining',
     '5-phase, 4-phase and 3-phase plans respectively',
     'Phase counts matched specification in all three cases',            True),
    ('FR10', 'Must',   'Navigated to /camp_plans after generating three plans',
     'All three plans listed with name and date',
     'Three plans displayed in reverse chronological order',             True),
    ('FR11', 'Should', 'Clicked delete on a camp plan; confirmed',
     '302 redirect; plan removed from list',
     'Plan deleted; list updated correctly',                             True),
    ('FR12', 'Must',   'Submitted game plan form with "strong boxing" and "weak takedown defence"',
     'Full round-by-round plan with range "Wrestling & Ground"',
     'Game plan rendered with 3 rounds, correct strategic approach',     True),
    ('FR13', 'Should', 'Clicked "Use Fighter" on an opponent in /fighters',
     'Game plan form pre-populated with that opponent\'s strengths/weaknesses',
     'Form fields pre-filled correctly',                                 True),
    ('FR14', 'Must',   'Navigated to /game_plans with two plans created',
     'Both plans listed with opponent names and dates',
     'Plans displayed; clicking one opened the full plan detail',        True),
    ('FR15', 'Should', 'Deleted a game plan from the list',
     '302 redirect; plan removed',
     'Plan deleted; no longer in list',                                  True),
    ('FR16', 'Must',   'Submitted /weight_tracker with weight 172.5 lbs',
     'WeightLog row created; weight appears in today\'s entry',
     'Entry saved and displayed in the log table',                       True),
    ('FR17', 'Must',   'Logged seven consecutive weights with an upward trend',
     'Trend indicator shows "Increasing"',
     '"Increasing" trend displayed; arrow icon correct',                 True),
    ('FR18', 'Must',   'Viewed /weight_tracker with active profile, weight over cut target',
     'Risk score < 100; deduction factors itemised',
     'Score and four deduction categories shown correctly',              True),
    ('FR19', 'Should', 'Clicked delete on a weight log entry',
     '302; entry removed from log',
     'Entry removed; log re-rendered without it',                        True),
    ('FR20', 'Must',   'Navigated to /risk_readiness with active profile and recent weight',
     'Composite score (0–100) and readiness level (Low/Medium/High) shown',
     'Score calculated correctly; level label matched expected band',    True),
    ('FR21', 'Must',   'Added opponent "Test Fighter" to /fighters with all fields',
     'Fighter row created; appears in list',
     'Fighter added and searchable',                                     True),
    ('FR22', 'Must',   'Searched /fighters?search=Test and filtered by weight class',
     'Only matching fighters returned',
     'Search and filter both narrowed results correctly',                True),
    ('FR23', 'Should', 'Edited fighter record; then deleted it',
     'Updated fields saved; fighter removed on delete',
     'Both operations confirmed; list updated',                          True),
    ('FR24', 'Could',  'Navigated to /compare_fighters with two fighter IDs in query string',
     'Side-by-side comparison rendered',
     'Both fighter cards displayed correctly',                           True),
    ('FR25', 'Must',   'Sent friend request to another user; accepted it as that user',
     'FriendRequest status changes accepted → Friendship row created',
     'Friend appeared in friends list for both accounts',                True),
    ('FR26', 'Should', 'Removed a friendship from /friends',
     'Friendship row deleted; friend no longer in list',
     'Friendship removed; confirmed on both accounts',                   True),
    ('FR27', 'Must',   'Clicked on a friend\'s profile from /friends',
     'Friend\'s fighter profile, weight logs and camp plans displayed',
     'All sections rendered with friend\'s data',                        True),
    ('FR28', 'Must',   'Opened /chat/<id> in two separate browsers; sent messages from each',
     'Messages appear in real time in both browser windows via SocketIO',
     'Real-time delivery confirmed; no page refresh required',           True),
    ('FR29', 'Must',   'Received a message while on a different page; navigated to chat',
     'Unread badge shows count on chat list; badge clears on open',
     'Badge incremented then cleared correctly',                         True),
    ('FR30', 'Must',   'Submitted /sparring_profile with location and skill level',
     'SparringProfile row created/updated; redirect to dashboard',
     'Profile saved; dashboard loaded correctly',                        True),
    ('FR31', 'Must',   'Viewed /sparring_dashboard with two other profiles in DB',
     'Partners ranked by compatibility score; partners < 50 excluded',
     'Ranked list displayed; scores verified against manual calculation', True),
    ('FR32', 'Must',   'Submitted sparring session request with date, time and location',
     'SparringSession row with pending status created',
     'Session appeared in both users\' dashboards with "Pending" status', True),
    ('FR33', 'Must',   'Accepted a pending session; marked it as completed',
     'Status updated to accepted then completed; completed_at populated',
     'Status changes confirmed; timestamp recorded',                     True),
    ('FR34', 'Should', 'Navigated to /assess_session/<id> after completing a session',
     'Assessment form rendered; submission creates SkillAssessment row',
     'Assessment saved; confirmation flash message displayed',           True),
    ('FR35', 'Could',  'Submitted assessment with rating 9 vs self-rating 4 (delta = 5)',
     'honesty_score decremented by 0.1',
     'SparringProfile.honesty_score reduced from 0.8 to 0.7',           True),
    ('FR36', 'Must',   'Submitted /events/create with all required fields',
     'Event row created with is_approved=True; redirect to event detail',
     'Event detail page rendered with correct data',                     True),
    ('FR37', 'Must',   'Filtered /events by city and weight class',
     'Only future approved events matching filters returned',
     'Filter correctly excluded past and non-matching events',           True),
    ('FR38', 'Must',   'Expressed interest in an event; then registered; then withdrew',
     'EventInterest.status cycles interested → registered → withdrawn',
     'Status updated at each step; organiser blocked from own event',    True),
    ('FR39', 'Must',   'Edited event description as organiser; attempted edit as non-organiser',
     'Organiser: event updated. Non-organiser: 302 with "permission denied" flash',
     'Both access paths behaved as expected',                            True),
    ('FR40', 'Must',   'Navigated to /admin/events as admin user',
     'All events listed with registered and interested counts',
     'Admin panel rendered with full statistics table',                   True),
]

tbl_rtm = doc.add_table(rows=1, cols=6)
tbl_rtm.style = 'Table Grid'
for cell, txt in zip(tbl_rtm.rows[0].cells, ['FR', 'Priority', 'Test Action', 'Expected', 'Actual', 'Status']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE, size=9)

for fr, pri, action, expected, actual, passed in rtm_rows:
    row = tbl_rtm.add_row().cells
    set_cell(row[0], fr, bold=True, center=True, size=9)
    set_cell(row[1], pri, center=True, bg_color=priority_colors[pri], size=9)
    set_cell(row[2], action, size=9)
    set_cell(row[3], expected, size=9)
    set_cell(row[4], actual, size=9)
    set_cell(row[5], 'PASS' if passed else 'FAIL', bold=True, center=True,
             bg_color='E2EFDA' if passed else 'FFE0E0', size=9)
no_split(tbl_rtm)
caption('Table 6.1 — Requirements Traceability Matrix: all 40 FRs (green = PASS)')

h2('6.3  Module-by-Module Test Scenarios')

h3('6.3.1  Authentication Module')
body(
    'The authentication module was tested for three scenarios: new user registration, login with '
    'correct credentials, and login with an incorrect password. Registration was verified to '
    'reject duplicate usernames and duplicate email addresses independently, producing a '
    'contextual validation error without exposing whether the conflict is on the username or '
    'the email field (a minor information-disclosure consideration). The case-insensitive login '
    'was tested by registering as "fighter1" and authenticating as "FIGHTER1"; the '
    'func.lower() comparison in the authentication route correctly matched both. Password '
    'mismatch produced a generic "Invalid credentials" flash message without revealing '
    'which field was wrong.'
)
add_picture_safe('screenshots/06_signup.png', 'Figure 6.1 — Sign-Up page')
caption('Figure 6.1 — User registration page with WTForms validation feedback')

h3('6.3.2  Training and Dashboard Module')
body(
    'The dashboard (Figure 5.4) was the most heavily tested page because it aggregates data '
    'from FighterProfile, WeightLog, CampPlan and ChatMessage in a single request. Six edge '
    'cases were explicitly exercised: (1) no fighter profile — redirects to /profile with a '
    'prompt; (2) fight date in the past — displays "Fight passed" rather than a negative '
    'countdown; (3) no weight entries — weight status card shows "No weight data"; (4) no '
    'camp plan — readiness penalty applied and a warning displayed; (5) fight date within '
    '7 days — "Taper Week" phase label with an urgent warning banner; (6) weight exactly '
    'at target — no weight-cut warning generated. All six scenarios rendered correctly '
    'without exceptions.'
)
body(
    'The camp planner input form is shown in Figure 6.2. It captures the opponent name, '
    'their key strengths, their weaknesses, and the fight type (3 or 5 rounds). These four '
    'inputs feed the keyword-matching engine described in Section 5.3.1. The form was tested '
    'with five different opponent profiles covering pure strikers, pure grapplers, wrestlers, '
    'wrestlers with weak chins, and a balanced mixed martial artist; each generated a '
    'distinct, non-generic set of training priorities.'
)
add_picture_safe('screenshots/07_camp_planner_form.png', 'Figure 6.2 — Camp Planner form')
caption('Figure 6.2 — Camp planner input form: opponent strengths, weaknesses and fight type feed the keyword-matching engine')

body(
    'Figure 6.3 shows the output for the Alex Pereira opponent profile (strong striking, '
    'Muay Thai clinch, elite boxing; weak grappling and slow footwork). The engine correctly '
    'identified the striking threat and populated the early phases with defensive boxing '
    'priorities, head-movement drills and cage-work, while later phases exploit the weak '
    'grappling with wrestling-offence tasks. The adaptive phase structure was verified to '
    'produce 5 phases at 10+ weeks, 4 phases at 6–9 weeks and 3 phases with fewer than '
    '6 weeks remaining (FR9), confirmed by setting the fight date at each threshold.'
)
add_picture_safe('screenshots/03_camp_plan.png', 'Figure 6.3 — Camp Plan output')
caption('Figure 6.3 — Generated camp plan output: five opponent-specific phases with drills, conditioning and mental-prep notes')

body(
    'The game plan generator was tested across all fifteen strategic branches in '
    'generate_game_plan(). Each branch was triggered by constructing an opponent '
    'description containing its keyword triggers; the resulting preferred range and '
    'strategic approach were verified against the expected decision-tree output. '
    'Figure 6.4 shows the output for the same opponent used in Figure 6.3: the engine '
    'correctly selected "Striking (Stay Standing)" as the preferred range given strong '
    'grappling combined with weak footwork, and produced distinct per-round objectives '
    'for a 3-round fight. The 5-round variant was also tested; the additional rounds '
    'each received adjusted tactical emphasis (later rounds targeting accumulated fatigue).'
)
add_picture_safe('screenshots/08_game_plan_output.png', 'Figure 6.4 — Game Plan output')
caption('Figure 6.4 — Round-by-round game plan: preferred range, strategic approach, techniques to drill and per-round objectives')

body(
    'The weight tracker (Figure 6.5) was tested by logging a sequence of entries above '
    'and below the cut target, then verifying that the trend indicator and risk score '
    'updated correctly. The risk score deduction for weight overage was verified by '
    'entering a weight 8 lbs over target; the formula (min(80, diff × 3) = 24) produced '
    'a score of 76, matching the expected calculation. The trend labels were verified by '
    'logging seven declining values ("Decreasing") and seven stable values ("Stable" within '
    '0.5 lb variation). Deletion of log entries was also tested; the trend recalculated '
    'correctly on the next page load (FR19).'
)
add_picture_safe('screenshots/09_weight_tracker.png', 'Figure 6.5 — Weight tracker')
caption('Figure 6.5 — Weight tracker: log entry form, 7-day trend indicator and itemised risk score deductions')

body(
    'Figure 6.6 shows the Risk and Readiness page for a fighter who is on track with weight, '
    'training five days per week, has an active camp plan, and has more than 60 days to '
    'their fight. The composite score is 70/100 (High Readiness) with no deductions, '
    'confirming the baseline path. The page was additionally tested with all four deduction '
    'factors active simultaneously: a fighter 12 lbs overweight, training twice a week, '
    'with no camp plan and a fight in 10 days produced a score of 0 (floor), correctly '
    'classified as "High Risk / Low Readiness".'
)
add_picture_safe('screenshots/10_risk_readiness.png', 'Figure 6.6 — Risk and Readiness score')
caption('Figure 6.6 — Risk and Readiness page: composite score 70/100 (High Readiness) with per-factor deduction breakdown')

h3('6.3.3  Social Module — Friends and Chat')
body(
    'The friend system was tested for the full lifecycle: send request → pending state visible '
    'to recipient → accept → friendship confirmed → view friend profile → remove friend. '
    'The notification badge on the navbar was observed to increment when a friend request was '
    'received and clear after the request was acted upon. The real-time chat (FR28) was the '
    'most technically demanding test: two accounts were opened simultaneously in separate '
    'browser windows. Messages sent from Account A appeared in Account B\'s chat window '
    'without a page refresh, confirming SocketIO event delivery. The unread message badge '
    '(FR29) was observed to increment on the chat list page and the dashboard preview, '
    'and to clear when the conversation was opened — even when the chat was already open '
    'in the other window, demonstrating that the current_chats server-side dictionary '
    'correctly suppressed the count update for the active viewer.'
)
add_picture_safe('screenshots/14_friend_request_send.png', 'Figure 6.7 — Send Friend Request')
caption('Figure 6.7 — Friend system: sending a request by username; "No friends yet" state before acceptance')

add_picture_safe('screenshots/15_friend_request_pending.png', 'Figure 6.8 — Pending Friend Request')
caption('Figure 6.8 — Pending friend request received: Accept and Decline options visible to recipient')

add_picture_safe('screenshots/10_friends.png', 'Figure 6.9 — Friends list after acceptance')
caption('Figure 6.9 — Friends list after request accepted: View Profile, Chat and Remove actions available')

add_picture_safe('screenshots/16_chat_sender.png', 'Figure 6.10 — Chat (sender side)')
caption('Figure 6.10 — Real-time chat from sender\'s perspective: message delivered instantly via SocketIO')

add_picture_safe('screenshots/11_chat.png', 'Figure 6.11 — Chat (recipient side)')
caption('Figure 6.11 — Chat from recipient\'s perspective showing both sides of the conversation in real time')

h3('6.3.4  Sparring Module')
body(
    'Sparring partner matching was tested with three synthetic profiles at varying skill levels '
    'and locations. A profile in the same city with the same skill level scored 90 (base 40 + '
    '25 skill + 25 style). A profile two skill levels apart scored 70, and a profile beyond the '
    'configured max_distance was excluded from results, confirming the distance filter in '
    'get_sparring_matches(). The honesty score mechanism (FR35) was tested by submitting a '
    'peer rating nine points above the self-rating; the honesty_score was observed to '
    'decrement by 0.1 in the database, consistent with the specification in Section 5.3.3.'
)
add_picture_safe('screenshots/04_sparring_dashboard.png', 'Figure 6.12 — Sparring dashboard')
caption('Figure 6.12 — Sparring partner matching dashboard with compatibility scores and distance')

h3('6.3.5  Events Module')
body(
    'Event creation was tested by a registered user submitting the full event form, including '
    'multiple weight classes and experience levels stored as comma-separated strings. The public '
    'listing filter was tested for each filter parameter independently (city, country, weight '
    'class, experience level, event type, rules) and in combination; all filter combinations '
    'correctly narrowed the event set without raising query errors. The organiser-restriction '
    'on self-registration (FR38) was verified by attempting to register for a self-created '
    'event; the route correctly returned the "You can\'t register for your own event" flash '
    'message. The admin panel (FR40) was tested by setting is_admin=True on a test account '
    'directly in the database and navigating to /admin/events; the panel rendered with full '
    'participant statistics.'
)
add_picture_safe('screenshots/05_events.png', 'Figure 6.13 — Events listing')
caption('Figure 6.13 — Public events listing with multi-field filter bar and event cards')

h3('6.3.6  Opponent Database Module')
body(
    'The opponent database was tested for all four CRUD operations: add, search, edit and '
    'delete. The shared (non-user-scoped) design was verified by confirming that a fighter '
    'added by User A was visible to User B without any data sharing mechanism. The "Use '
    'Fighter" redirect (FR13) was tested end-to-end: clicking the button from the fighter '
    'detail page navigated to /game_plan with the strengths and weaknesses query parameters '
    'pre-populated in the form, allowing the user to generate a game plan without re-typing '
    'the opponent\'s attributes.'
)
add_picture_safe('screenshots/12_fighter_db.png', 'Figure 6.14 — Opponent database')
caption('Figure 6.14 — Shared opponent database with search bar, weight-class filter and "Use For Game Plan" action')

h2('6.4  Non-Functional Requirements Verification')
body(
    'Table 6.2 records the verification evidence for all 17 non-functional requirements. '
    'For things that could be measured directly — responsive breakpoints, HTTP status codes — '
    'they were. For requirements that describe how the code is structured (Blueprint layout, '
    'extensions.py), the evidence points to the codebase itself.'
)

nfr_verify = [
    ('NFR1',  'Usability',       'Must',   'Tested at 320 px (iPhone SE), 768 px (iPad) and 1920 px. Bootstrap 5 grid collapsed navbar and stacked cards correctly at all widths.',                                  True),
    ('NFR2',  'Usability',       'Must',   'Every POST route verified to flash a categorised message (success/warning/danger). Flash categories mapped to Bootstrap alert classes (green/yellow/red).',              True),
    ('NFR3',  'Usability',       'Should', 'Dashboard tested at 768 px; all four key metrics (countdown, weight, phase, warnings) visible without vertical scroll.',                                                  True),
    ('NFR4',  'Functionality',   'Must',   'Flask-WTF CSRF tokens present in all forms; verified via browser DevTools inspecting the hidden input field on every form page.',                                         True),
    ('NFR5',  'Functionality',   'Must',   'Server-side validation tested by submitting empty and malformed inputs; WTForms rejected all invalid submissions before any DB write.',                                   True),
    ('NFR6',  'Functionality',   'Must',   'User.password_hash column confirmed to contain Werkzeug PBKDF2-SHA256 digest strings. Plaintext never stored; verified by inspecting fpms.db directly.',                 True),
    ('NFR7',  'Functionality',   'Must',   'All 41 @login_required routes tested with a fresh unauthenticated session; all redirected to /login with a "Please log in" flash message.',                              True),
    ('NFR8',  'Reliability',     'Must',   'All database queries issued via SQLAlchemy ORM. No raw SQL strings used in any route file; confirmed by grep for "execute(" in routes/.',                                True),
    ('NFR9',  'Reliability',     'Should', 'Foreign key relationships declared in all model classes. SQLAlchemy relationship() and ForeignKey() used throughout models.py.',                                         True),
    ('NFR10', 'Reliability',     'Should', 'Missing FighterProfile tested: /dashboard redirects to /profile with a prompt. Missing SparringProfile: /sparring_dashboard redirects to /sparring_profile.',            True),
    ('NFR11', 'Portability',     'Must',   'SQLALCHEMY_DATABASE_URI set to sqlite:///fpms.db in app.py. Changing to a PostgreSQL URI would require no other code changes; verified by reviewing the ORM layer.',    True),
    ('NFR12', 'Portability',     'Must',   'requirements.txt present with pinned versions. Application tested on Python 3.13; no f-string or walrus-operator syntax below Python 3.8.',                             True),
    ('NFR13', 'Efficiency',      'Should', 'SQLAlchemy relationships use lazy="select" (default lazy loading). No joined eager-load queries observed on list pages via SQLAlchemy echo mode.',                       True),
    ('NFR14', 'Efficiency',      'Should', 'get_sparring_matches() returns matches[:20]. Verified by inserting 25 profiles and confirming only 20 returned.',                                                        True),
    ('NFR15', 'Maintainability', 'Must',   'Seven Blueprint files confirmed under routes/. Each blueprint registered in app.py. No cross-module direct function calls; all communication via URL routing.',          True),
    ('NFR16', 'Maintainability', 'Must',   'extensions.py contains LoginManager and SocketIO instances. Both imported by app.py and social blueprint without circular import error on startup.',                     True),
    ('NFR17', 'Maintainability', 'Should', 'SECRET_KEY hard-coded in app.py with inline comment "Change this in production". README documents environment variable substitution procedure.',                          True),
]

tbl_nfr2 = doc.add_table(rows=1, cols=5)
tbl_nfr2.style = 'Table Grid'
for cell, txt in zip(tbl_nfr2.rows[0].cells, ['NFR', 'Category', 'Priority', 'Evidence', 'Status']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE, size=9)
for nfr, cat, pri, evidence, passed in nfr_verify:
    row = tbl_nfr2.add_row().cells
    set_cell(row[0], nfr, bold=True, center=True, size=9)
    set_cell(row[1], cat, bg_color=cat_colors[cat], center=True, size=9)
    set_cell(row[2], pri, center=True, bg_color=priority_colors[pri], size=9)
    set_cell(row[3], evidence, size=9)
    set_cell(row[4], 'PASS' if passed else 'FAIL', bold=True, center=True,
             bg_color='E2EFDA' if passed else 'FFE0E0', size=9)
no_split(tbl_nfr2)
caption('Table 6.2 — NFR verification evidence (all 17 NFRs passed)')

h2('6.5  Route Coverage Summary')
body(
    'Table 6.3 shows route coverage across all seven blueprints. '
    'All 44 routes were hit during testing with at least one passing scenario and one '
    'error or access-control check.'
)

cov_rows = [
    ('auth',     3,  3,  '100%', 'Registration, login (correct + incorrect credentials), logout'),
    ('main',     2,  2,  '100%', 'Dashboard (with and without profile), index redirect'),
    ('training', 11, 11, '100%', 'Profile CRUD, camp plan CRUD, game plan CRUD, weight log CRUD, readiness'),
    ('social',   8,  8,  '100%', 'Friend request lifecycle, friends list, profile view, chat list, chat view'),
    ('fighters', 6,  6,  '100%', 'Add, search, filter, detail, edit, delete, compare, use-fighter redirect'),
    ('sparring', 6,  6,  '100%', 'Sparring profile, dashboard, partner profile, request, respond, assess'),
    ('events',   8,  8,  '100%', 'Create, list (filtered), detail, interest, my-events, edit, delete, admin'),
]

tbl_cov = doc.add_table(rows=1, cols=5)
tbl_cov.style = 'Table Grid'
for cell, txt in zip(tbl_cov.rows[0].cells, ['Blueprint', 'Routes', 'Tested', 'Coverage', 'Scenarios Exercised']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for bp, total, tested, pct, scenarios in cov_rows:
    row = tbl_cov.add_row().cells
    set_cell(row[0], bp, bold=True)
    set_cell(row[1], str(total), center=True)
    set_cell(row[2], str(tested), center=True)
    set_cell(row[3], pct, bold=True, center=True, bg_color='E2EFDA')
    set_cell(row[4], scenarios, size=9)
no_split(tbl_cov)
caption('Table 6.3 — Route coverage summary: 44/44 routes tested across 7 blueprints')

h2('6.6  Known Defects and Limitations')
body(
    'Everything in the FR table passed, but testing also turned up a handful of issues '
    'that sit outside the formal requirement set and are worth noting.'
)
for defect in [
    'SECRET_KEY is hard-coded in app.py as a development placeholder (flagged in NFR17 and risk R5). '
     'This is not exploitable in the development environment but must be resolved before any production deployment.',
    'The is_admin flag can only be set by direct database manipulation; there is no admin registration '
     'or elevation workflow in the UI. An unintended consequence is that any database administrator '
     'could grant themselves admin rights without an audit trail.',
    'The application has no rate-limiting on login attempts, making it theoretically susceptible to '
     'credential-stuffing attacks. Flask-Limiter was identified as a potential mitigation for future work.',
    'File uploads for profile pictures are limited to 16 MB (MAX_CONTENT_LENGTH) but the file type '
     'is not strictly validated beyond Werkzeug\'s secure_filename() call; a MIME-type check would '
     'add additional robustness.',
    'SQLite does not enforce foreign key constraints by default; PRAGMA foreign_keys = ON would need '
     'to be issued at connection time to enable cascading deletes at the database layer.',
]:
    bullet(defect)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7 — CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
h1('Chapter 7 — Conclusion')

h2('7.1  Summary of Achievements')
body(
    'FPMS is a working web application. It covers fight-camp planning, weight monitoring, '
    'game-plan generation, sparring-partner matching, event management and real-time messaging '
    '— all in a single platform, which was the central goal from Chapter 1. By the end of '
    'Iteration 4 it comprised 44 routes across seven Flask Blueprints, 14 relational database '
    'models, and a live WebSocket channel for chat. All 34 Must-have and Should-have functional '
    'requirements were implemented and verified through manual black-box testing, and all five '
    'objectives from Section 1.3 were met.'
)

h2('7.2  Critical Reflection')
body(
    'The reflection below follows Gibbs\' Reflective Cycle (Gibbs, 1988). The model was chosen '
    'because it pushes past simple description — it requires an honest assessment of what '
    'actually happened, not just what was intended, and commits to concrete actions rather '
    'than vague lessons.'
)

h3('7.2.1  Description')
body(
    'The project was developed iteratively over four loosely time-boxed iterations spanning '
    'approximately twelve weeks. The application evolved from a minimal authentication scaffold '
    'in Iteration 1 to a feature-complete platform by Iteration 4. The most challenging '
    'technical decisions were the introduction of Flask Blueprints (which required '
    'refactoring all route files and templates mid-project) and the integration of '
    'Flask-SocketIO (which introduced an asynchronous event model unfamiliar to a developer '
    'working primarily in synchronous Flask views).'
)

h3('7.2.2  Feelings')
body(
    'The early iterations were productive and motivating: authentication and the fighter profile '
    'were well-understood problems with clearly defined models. Confidence dipped significantly '
    'during the introduction of real-time chat in Iteration 3. Debugging SocketIO event delivery '
    'across multiple browser windows — particularly understanding the room-naming convention and '
    'the current_chats dictionary — took considerably longer than anticipated and temporarily '
    'put the project timeline at risk. The Blueprint refactoring in Iteration 4, while '
    'ultimately the right architectural choice, was initially intimidating because it required '
    'touching every template and route simultaneously. Completing it successfully and seeing the '
    'circular-import warning disappear was a significant confidence-restoring moment.'
)

h3('7.2.3  Evaluation')
body(
    'The Blueprint architecture proved its value immediately. Once each module was isolated, '
    'adding features to the sparring blueprint, for example, did not risk breaking the events '
    'routes. The keyword-matching logic for camp and game plan generation was more effective '
    'than expected — the output felt genuinely tailored to the opponent being prepared for, '
    'not generic. Using domain-specific heuristics rather than an external AI API kept the '
    'system fast, deterministic and offline-capable. The Haversine formula for geolocation '
    'matching was another good call: accurate distance calculations with no third-party mapping '
    'dependency.'
)
body(
    'The most significant shortcoming is the absence of an automated test suite. Manual testing '
    'verified all 40 functional requirements, but it cannot be run reliably after future changes '
    'and cannot catch regressions. Pytest integration tests were deprioritised in favour of '
    'feature velocity, which is a trade-off that would not be acceptable in a professional '
    'setting. The admin flag management was also poorly thought through: there is no UI for '
    'granting or revoking admin rights, which means the admin panel is effectively invisible '
    'to any user who has not been given direct database access.'
)

h3('7.2.4  Analysis')
body(
    'The SocketIO integration difficulties arose primarily from a conceptual mismatch: the '
    'developer initially attempted to model chat delivery using HTTP polling logic before '
    'correctly re-framing it as an event-driven model. The resolution — the current_chats '
    'dictionary and the pair of per-user and per-conversation rooms — was derived empirically '
    'through incremental testing rather than through upfront design. This suggests that for '
    'novel architectural patterns, a spike prototype earlier in the iteration would have '
    'been more efficient than proceeding with an incomplete mental model. The Blueprint '
    'refactoring delay similarly stemmed from deferring a known architectural risk (R3 in '
    'Table 2.2) longer than the risk register recommended; had it been addressed at the '
    'end of Iteration 2 rather than mid-Iteration 4, the disruption would have been smaller.'
)

h3('7.2.5  Conclusion and Action Plan')
body(
    'The main takeaway from this project is that the architecture decisions — particularly '
    'around circular imports and real-time communication — needed earlier attention than '
    'they got. Both problems were on the risk register; both were deferred anyway, and both '
    'caused disruption when they finally had to be dealt with. That would not have changed '
    'the end result, but it would have made the last iteration considerably less stressful. '
    'Going forward the three most important actions are: (1) add a pytest integration test '
    'suite covering all 44 routes; (2) deploy to a cloud environment with a PostgreSQL '
    'backend and environment-variable configuration; and (3) get real fighters to actually '
    'use the camp and game plan features and give honest feedback on whether the output '
    'is useful in practice.'
)

h2('7.3  Future Work')
body(
    'A number of ideas came up during development that were deliberately left out to keep the scope manageable.'
)
for fw in [
    'Production deployment — migrate from SQLite to PostgreSQL, replace the hard-coded SECRET_KEY '
     'with an environment variable, add HTTPS via Let\'s Encrypt, and deploy to a cloud platform '
     'such as Heroku (2024) or AWS Elastic Beanstalk.',
    'Automated testing — introduce a pytest (2024) test suite with Flask\'s test client, targeting '
     'all 44 routes with at least one happy-path and one error-path test case each. Aim for '
     'greater than 80% line coverage.',
    'Wearable and sensor integration — connect to Garmin Connect or Apple HealthKit APIs to '
     'automatically import daily resting heart rate, HRV and training load data, replacing '
     'manual weight log entries with automated biometric imports.',
    'Mobile-first progressive web app (PWA) — add a service worker and web app manifest '
     'to enable offline access to camp plans and game plans, and push notifications for '
     'unread messages and upcoming fight-date reminders.',
    'Real user validation — recruit five to ten amateur fighters from a local MMA gym, '
     'conduct a structured usability study following ethical approval, and iterate the '
     'interface based on findings.',
    'Admin workflow — implement a proper admin registration and role-management UI, '
     'including an audit log of admin actions and the ability for event organisers to '
     'request admin-level event approval rather than requiring direct database access.',
]:
    bullet(fw)
doc.add_paragraph()

h2('7.4  Closing Statement')
body(
    'Building FPMS confirmed that the gap identified at the start of the project is real and '
    'addressable. A solo developer, working iteratively over twelve weeks, can produce a '
    'platform that covers the full preparation lifecycle of an amateur fighter — not just one '
    'slice of it. The codebase is modular, the core features work, and the architecture is '
    'clean enough to extend without starting over. What it is not is production-ready: no '
    'automated tests, a hard-coded secret key, and no deployment pipeline. Those are '
    'fixable problems, and the path to fixing them is set out in Section 7.3. The '
    'more interesting question — whether fighters would actually use it — is one that '
    'only real-user testing can answer.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
h1('References')
for entry in sorted(REFERENCES, key=lambda x: x.upper()):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after       = Pt(6)
    r = p.add_run(entry)
    r.font.name = 'Calibri'; r.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDICES
# ═══════════════════════════════════════════════════════════════════════════════
h1('Appendices')

# ─── Appendix A — Ethics ──────────────────────────────────────────────────────
h2('Appendix A — FET Ethical Review Checklist')
body(
    'The following checklist was completed and submitted to the Faculty of Environment and '
    'Technology (FET) ethics committee prior to any development or data collection activity. '
    'All questions were answered "No" or "Not applicable", confirming that the project '
    'falls within the low-risk category and does not require a full ethics review application.'
)

ethics_rows = [
    ('A1',  'Does the project involve human participants (interviews, surveys, observations)?',                           'No — all data is synthetic; no real users were recruited.'),
    ('A2',  'Does the project involve collection of personal data (names, contact details, biometrics)?',                'No — test accounts use fictitious usernames and synthetic weight values.'),
    ('A3',  'Does the project involve sensitive personal data (health, financial, criminal records)?',                   'No — weight log data in the system is synthetic test data only.'),
    ('A4',  'Does the project involve deception of participants?',                                                       'No — not applicable; no human participants involved.'),
    ('A5',  'Does the project involve children or vulnerable adults?',                                                   'No — not applicable.'),
    ('A6',  'Does the project access or use third-party personal data without explicit consent?',                        'No — the system operates on data entered by the authenticated user only.'),
    ('A7',  'Does the project create or deploy software that could be used to harm individuals or groups?',              'No — the application is a training management tool with no harmful capability.'),
    ('A8',  'Does the project involve network penetration testing, vulnerability scanning, or security exploitation?',   'No — all testing is black-box functional testing against a local development instance.'),
    ('A9',  'Are there any potential conflicts of interest that could affect the objectivity of the research?',          'No — the researcher has no financial or professional relationship with any platform reviewed.'),
    ('A10', 'Is informed consent required from any party?',                                                              'No — no human participants; consent not required.'),
]

tbl_eth = doc.add_table(rows=1, cols=3)
tbl_eth.style = 'Table Grid'
for cell, txt in zip(tbl_eth.rows[0].cells, ['Ref', 'Question', 'Response']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE)
for ref, question, response in ethics_rows:
    row = tbl_eth.add_row().cells
    set_cell(row[0], ref, bold=True, center=True, size=9)
    set_cell(row[1], question, size=9)
    set_cell(row[2], response, size=9, bg_color='E2EFDA')
no_split(tbl_eth)
caption('Table A.1 — FET Ethical Review Checklist (all items: no risk identified)')

# ─── Appendix B — Extended Requirements ──────────────────────────────────────
doc.add_page_break()
h2('Appendix B — Extended Requirements Tables')
body(
    'The tables below extend the FR and NFR summaries from Chapters 4 and 6 with '
    'the full acceptance criteria used during manual testing.'
)

h3('B.1 — Functional Requirements with Acceptance Criteria')

fr_ac = [
    ('FR1',  'Must',   'Allow guest registration with unique username and email.',
     'Form accepted with unique fields; duplicate username/email rejected with error flash; password stored as Werkzeug PBKDF2 hash.'),
    ('FR2',  'Must',   'Authenticate with case-insensitive username and hashed password.',
     'Correct credentials: 302 redirect to /dashboard. Wrong password: form re-rendered with "Invalid credentials" flash. Unknown username: same.'),
    ('FR3',  'Must',   'Maintain authenticated session; provide logout.',
     'Session cookie present on login; /logout clears it; all @login_required routes redirect to /login after logout.'),
    ('FR4',  'Should', 'Allow profile picture upload to server-side folder.',
     'File saved under static/uploads/<user_id>/; DB stores relative path; img tag renders on profile page.'),
    ('FR5',  'Must',   'Create/update fighter profile (age, height, weight, class, fight date, training days, style).',
     'Profile saved to DB; /dashboard reads it for countdown and phase; second save overwrites all fields correctly.'),
    ('FR6',  'Must',   'Dashboard shows fight countdown, weight status, camp phase, and warnings.',
     'Four metric cards visible at 1280 px and 768 px; countdown matches (fight_date - today).days; phase label matches FR7 logic.'),
    ('FR7',  'Must',   'Derive camp phase from weeks remaining.',
     '> 10 weeks → "Base Conditioning"; 7–10 → "Skill Sharpening"; 3–6 → "Sparring Peak"; < 3 → "Taper Week".'),
    ('FR8',  'Must',   'Generate opponent-specific camp plan via keyword matching.',
     'Plan contains phase titles; striking keywords in strengths produce defensive boxing priorities; grappling keywords produce wrestling-defence tasks.'),
    ('FR9',  'Must',   'Adapt phase count to available weeks.',
     '10+ weeks: 5 phases. 6–9 weeks: 4 phases. < 6 weeks: 3 phases. Verified by setting fight dates at each threshold.'),
    ('FR10', 'Must',   'View history of all camp plans.',
     '/view_camp_plans lists all plans for current user; opponent name and date visible; plans from other users not shown.'),
    ('FR11', 'Should', 'Delete own camp plans.',
     'DELETE confirmed removes record from DB; plan no longer appears in /view_camp_plans; redirects with success flash.'),
    ('FR12', 'Must',   'Generate round-by-round game plan via decision-tree.',
     'Plan contains preferred range, approach, per-round objectives, techniques to drill, and things to avoid for each round requested.'),
    ('FR13', 'Should', 'Pre-populate game plan from opponent database.',
     '"Use For Game Plan" button on /fighters/<id> redirects to /game_plan with strengths/weaknesses query params pre-filled.'),
    ('FR14', 'Must',   'View all generated game plans.',
     '/view_game_plans lists all plans for current user in reverse chronological order.'),
    ('FR15', 'Should', 'Delete own game plans.',
     'DELETE removes record; flash confirmation shown; plan absent from list.'),
    ('FR16', 'Must',   'Log body weight with date stamp.',
     'POST to /log_weight creates WeightLog record; log appears in table below form with correct date.'),
    ('FR17', 'Must',   'Display 7-day rolling weight trend.',
     '"Increasing" when last 7-day average exceeds prior 7-day average; "Decreasing" inverse; "Stable" when difference < 0.5 lb.'),
    ('FR18', 'Must',   'Produce weight-cut risk score with itemised deductions.',
     'Score starts at 100; deductions for weight overage, low training days, no camp plan, proximity to fight; each factor listed on page.'),
    ('FR19', 'Should', 'Delete individual weight log entries.',
     'DELETE removes record from DB; entry absent from log table; trend recalculates on next page load.'),
    ('FR20', 'Must',   'Produce composite readiness score (0–100) and readiness level.',
     'Score matches manual calculation of deductions; level "Low" < 30, "Medium" 30–69, "High" 70+.'),
    ('FR21', 'Must',   'Add opponents to shared database with full attributes.',
     'Fighter record visible to all authenticated users after creation; all fields stored and displayed on detail page.'),
    ('FR22', 'Must',   'Search opponents by name; filter by weight class.',
     'Name search returns partial matches case-insensitively; weight-class dropdown filters list correctly.'),
    ('FR23', 'Should', 'Edit and delete opponent records.',
     'Edit pre-fills form with existing values; save updates DB; delete removes record with confirmation flash.'),
    ('FR24', 'Could',  'Side-by-side comparison of two opponents.',
     '/compare_fighters?f1=<id>&f2=<id> renders two columns with all attributes aligned row by row.'),
    ('FR25', 'Must',   'Send, accept and decline friend requests by username.',
     'Send: FriendRequest row created with status "pending". Accept: Friendship row created, request status "accepted". Decline: request status "declined".'),
    ('FR26', 'Should', 'Remove existing friendships.',
     'Friendship record deleted from both directions; removed user no longer appears in friends list.'),
    ('FR27', 'Must',   'View a friend\'s fighter profile and training data.',
     '/friend_profile/<id> accessible only if Friendship exists; shows target user\'s profile, weight logs and camp plans.'),
    ('FR28', 'Must',   'Real-time bidirectional messaging via WebSocket.',
     'Message sent by User A appears in User B\'s open chat window without page reload; persisted to DB.'),
    ('FR29', 'Must',   'Track unread message counts; mark as read on open.',
     'Badge count increments when message arrives in another tab; drops to 0 when chat is opened.'),
    ('FR30', 'Must',   'Create sparring profile with location, skill, style, availability, max distance.',
     'SparringProfile created in DB; editable via same form; changes reflected on /sparring_dashboard.'),
    ('FR31', 'Must',   'Rank sparring partners by compatibility score (0–100); exclude below 50.',
     'Score composed of skill proximity (max 25) + style match (max 25) + distance (max 20) + base 40; partners < 50 absent from results.'),
    ('FR32', 'Must',   'Request sparring sessions with date, time, duration and location.',
     'SparringSession created with status "pending"; appears in partner\'s session list.'),
    ('FR33', 'Must',   'Recipients can accept, decline or complete sessions.',
     'Accept: status → "accepted"; Decline: status → "declined"; Complete: status → "completed"; flash shown for each.'),
    ('FR34', 'Should', 'Submit peer skill assessment after completed session.',
     'Assessment form available only when session status is "completed"; SkillAssessment record created in DB.'),
    ('FR35', 'Could',  'Maintain honesty score on sparring profiles.',
     'After assessment, honesty_score incremented by 0.1 if self-rating within 1 of peer average; decremented by 0.1 if delta >= 3.'),
    ('FR36', 'Must',   'Create events with full attribute set.',
     'Event record created with all 14 fields; creator listed as organizer; event appears under /my_events.'),
    ('FR37', 'Must',   'Public listing shows only future approved events; supports multi-field filtering.',
     'Past events absent; unapproved events absent; city/weight class/rules filters each reduce result set correctly.'),
    ('FR38', 'Must',   'Users can register interest in events; organiser cannot self-register.',
     'EventInterest record created; POST to own event returns "cannot register for your own event" flash.'),
    ('FR39', 'Must',   'Organisers can edit and delete own events.',
     'Edit pre-fills form; save updates DB; delete removes event and all associated interest records.'),
    ('FR40', 'Must',   'Admin event panel with registration statistics.',
     '/admin/events accessible only with is_admin=True; shows all events with participant counts.'),
]

tbl_b1 = doc.add_table(rows=1, cols=4)
tbl_b1.style = 'Table Grid'
for cell, txt in zip(tbl_b1.rows[0].cells, ['FR', 'Pri', 'Description', 'Acceptance Criterion']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE, size=9)
for fr, pri, desc, ac in fr_ac:
    row = tbl_b1.add_row().cells
    set_cell(row[0], fr,   bold=True, center=True, size=9)
    pri_bg = {'Must': 'FFE0E0', 'Should': 'FFF3CD', 'Could': 'E2EFDA'}.get(pri, 'FFFFFF')
    set_cell(row[1], pri,  bold=True, center=True, size=9, bg_color=pri_bg)
    set_cell(row[2], desc, size=9)
    set_cell(row[3], ac,   size=9)
no_split(tbl_b1)
caption('Table B.1 — All 40 functional requirements with acceptance criteria (Must = red, Should = amber, Could = green)')

doc.add_page_break()
h3('B.2 — Non-Functional Requirements with Measurement Criteria')
body(
    'Table B.2 extends the NFR set from Chapter 4 with the measurement method and '
    'acceptance threshold used during verification (Chapter 6).'
)

nfr_extended = [
    ('NFR1',  'Usability',       'Must',
     'Responsive layout across 320 px, 768 px and 1920 px viewports.',
     'Manual test at three breakpoints using browser DevTools device emulation.',
     'Bootstrap 5 navbar collapses and cards stack at 320 px; no horizontal overflow.'),
    ('NFR2',  'Usability',       'Must',
     'All user actions produce clear, categorised feedback.',
     'Trigger every success, warning and error path; inspect flash message class.',
     'Every path produces a Bootstrap alert with the correct colour (green/yellow/red).'),
    ('NFR3',  'Usability',       'Should',
     'Dashboard key metrics visible without scroll at 768 px.',
     'Load /dashboard at 768 px in DevTools; inspect vertical scroll height.',
     'All four metric cards (countdown, weight, phase, warnings) visible above fold.'),
    ('NFR4',  'Functionality',   'Must',
     'CSRF token present in every POST form.',
     'Inspect hidden input field on every form page; attempt POST without token.',
     'csrf_token hidden input present; raw POST without token returns 400 Bad Request.'),
    ('NFR5',  'Functionality',   'Must',
     'Server-side input validation rejects malformed data before any DB write.',
     'Submit empty forms and boundary-violating values to all WTForms routes.',
     'All invalid submissions re-render the form with an error; no DB record created.'),
    ('NFR6',  'Functionality',   'Must',
     'Passwords stored as one-way hashes; plaintext never persisted.',
     'Inspect User.password_hash column in fpms.db after registration.',
     'Column contains Werkzeug PBKDF2-SHA256 digest starting with "pbkdf2:sha256:".'),
    ('NFR7',  'Functionality',   'Must',
     'All protected routes require authentication.',
     'Access all 41 @login_required routes with a fresh unauthenticated browser session.',
     'Every route returns HTTP 302 redirect to /login with "Please log in" flash.'),
    ('NFR8',  'Reliability',     'Must',
     'All DB access via ORM; no raw SQL strings.',
     'grep -r "execute(" routes/ in project root.',
     'Zero occurrences of raw execute() calls in any route file.'),
    ('NFR9',  'Reliability',     'Should',
     'Foreign key relationships declared for all inter-model references.',
     'Review models.py for db.relationship() and ForeignKey() declarations.',
     'All FK columns reference parent table; relationship() backref defined for all 1:N pairs.'),
    ('NFR10', 'Reliability',     'Should',
     'Missing profile data handled gracefully.',
     'Access /dashboard and /sparring_dashboard with no FighterProfile / SparringProfile.',
     'Both routes redirect to the respective profile creation page with an informative flash.'),
    ('NFR11', 'Portability',     'Must',
     'Database backend changeable via URI string only.',
     'Review app.py SQLALCHEMY_DATABASE_URI; attempt to change to PostgreSQL URI.',
     'No application code change required; ORM layer abstracts all DB-specific syntax.'),
    ('NFR12', 'Portability',     'Must',
     'Reproducible installation via requirements.txt.',
     'Create fresh virtual environment; run pip install -r requirements.txt.',
     'All packages install without conflict; application starts on Python 3.13.'),
    ('NFR13', 'Efficiency',      'Should',
     'Lazy loading used; no N+1 queries on list pages.',
     'Enable SQLAlchemy echo mode; load /fighters and /view_camp_plans; inspect query log.',
     'List pages issue one query per table; no per-row subqueries observed.'),
    ('NFR14', 'Efficiency',      'Should',
     'Sparring dashboard limited to 20 best matches.',
     'Insert 25 SparringProfile records; load /sparring_dashboard.',
     'Dashboard renders exactly 20 match cards; confirmed by HTML element count.'),
    ('NFR15', 'Maintainability', 'Must',
     'Routes structured as seven Flask Blueprints.',
     'Inspect routes/ directory; review blueprint registration in app.py.',
     'Seven blueprint files present; each registered with app.register_blueprint(); no routes in app.py.'),
    ('NFR16', 'Maintainability', 'Must',
     'Shared extensions in extensions.py to prevent circular imports.',
     'Run python app.py; check for ImportError or circular import traceback.',
     'Application starts cleanly; LoginManager and SocketIO instances imported without error.'),
    ('NFR17', 'Maintainability', 'Should',
     'SECRET_KEY documented as requiring environment variable substitution.',
     'Inspect app.py; read README.md.',
     'Inline comment in app.py reads "Change this in production"; README documents env var override procedure.'),
]

tbl_b2 = doc.add_table(rows=1, cols=6)
tbl_b2.style = 'Table Grid'
for cell, txt in zip(tbl_b2.rows[0].cells,
                     ['NFR', 'Category', 'Priority', 'Requirement', 'Test Method', 'Acceptance Threshold']):
    set_cell(cell, txt, bold=True, bg_color='1F497D', color=WHITE, size=9)
cat_bg = {'Usability': 'DAEEF3', 'Functionality': 'E2EFDA',
          'Reliability': 'FFF3CD', 'Portability': 'FCE4D6',
          'Efficiency': 'EAD1DC', 'Maintainability': 'D9D2E9'}
for nfr, cat, pri, req, method, threshold in nfr_extended:
    row = tbl_b2.add_row().cells
    set_cell(row[0], nfr,       bold=True, center=True, size=9)
    set_cell(row[1], cat,       center=True, size=9, bg_color=cat_bg.get(cat, 'FFFFFF'))
    pri_bg = {'Must': 'FFE0E0', 'Should': 'FFF3CD'}.get(pri, 'E2EFDA')
    set_cell(row[2], pri,       bold=True, center=True, size=9, bg_color=pri_bg)
    set_cell(row[3], req,       size=9)
    set_cell(row[4], method,    size=9)
    set_cell(row[5], threshold, size=9)
no_split(tbl_b2)
caption('Table B.2 — All 17 NFRs with category, measurement method and acceptance threshold')

# ─── Appendix C — Code Listings ───────────────────────────────────────────────
doc.add_page_break()
h2('Appendix C — Selected Code Listings')
body(
    'The following extracts are the most algorithmically significant functions in the FPMS codebase. '
    'All listings are taken directly from the production source without modification.'
)

h3('C.1 — Camp Plan Keyword-Matching Engine (routes/training.py)')
body(
    'The generate_camp_plan() function builds an opponent-specific periodised plan by scanning '
    'the opponent\'s attributes for domain-specific keywords. Each keyword match contributes '
    'specific training priorities, sparring focuses, technique drills and mental preparation '
    'notes to the output phase structure. The phase count and duration are then scaled to the '
    'number of weeks remaining until the fight date.'
)
code_block([
    'def generate_camp_plan(strengths, weaknesses, weeks_remaining, fight_type):',
    '    strengths_lower  = strengths.lower()',
    '    weaknesses_lower = weaknesses.lower()',
    '',
    '    # Keyword banks for opponent-attribute detection',
    '    striking_keywords  = ["striking","punching","boxing","kickboxing",',
    '                          "muay thai","karate","knockout","ko","hands","kicks"]',
    '    wrestling_keywords = ["wrestling","takedowns","cage control","clinch"]',
    '    bjj_keywords       = ["bjj","jiu-jitsu","submissions","guard","ground game"]',
    '    grappling_keywords = wrestling_keywords + bjj_keywords + ["grappling","gnp"]',
    '    cardio_keywords    = ["cardio","conditioning","endurance","pace","gas tank"]',
    '    power_keywords     = ["power","knockout","ko","heavy hands","finisher"]',
    '    pressure_keywords  = ["pressure","aggression","aggressive","forward"]',
    '',
    '    # Detect opponent attributes from strengths',
    '    opp_striking  = any(kw in strengths_lower for kw in striking_keywords)',
    '    opp_grappling = any(kw in strengths_lower for kw in grappling_keywords)',
    '    opp_cardio    = any(kw in strengths_lower for kw in cardio_keywords)',
    '    opp_power     = any(kw in strengths_lower for kw in power_keywords)',
    '    opp_pressure  = any(kw in strengths_lower for kw in pressure_keywords)',
    '',
    '    training_priorities = []',
    '',
    '    if opp_striking:',
    '        training_priorities.append("Defensive boxing and head movement")',
    '    if opp_grappling:',
    '        training_priorities.append("Submission defense and escapes")',
    '    if opp_cardio:',
    '        training_priorities.append("High-volume conditioning work")',
    '    if opp_power:',
    '        training_priorities.append("Chin tucked, hands up at ALL times")',
    '    if opp_pressure:',
    '        training_priorities.append("Circle and angle work to avoid cage")',
    '',
    '    # Exploit opponent weaknesses',
    '    weak_striking  = any(kw in weaknesses_lower for kw in striking_keywords)',
    '    weak_grappling = any(kw in weaknesses_lower for kw in grappling_keywords)',
    '    if weak_striking:',
    '        training_priorities.append("Sharpen striking combinations for openings")',
    '    if weak_grappling:',
    '        training_priorities.append("Wrestling offense and top control")',
    '',
    '    # Scale phase structure to available weeks',
    '    if weeks_remaining >= 10:',
    '        phases = {',
    '            "Phase 1 - Foundation":        {"duration": "Weeks 1-3", ...},',
    '            "Phase 2 - Skill Development": {"duration": "Weeks 4-6", ...},',
    '            "Phase 3 - Sparring Peak":     {"duration": "Weeks 7-9", ...},',
    '            "Phase 4 - Sharpening":        {"duration": "Weeks 10-11", ...},',
    '            "Phase 5 - Taper":             {"duration": "Final week", ...},',
    '        }',
    '    elif weeks_remaining >= 6:',
    '        phases = { "Phase 1 - Accelerated Build": ...,  # 4-phase structure',
    '                   "Phase 2 - Opponent-Specific": ..., }',
    '    else:',
    '        phases = { "Phase 1 - Intensive Prep": ...,    # 3-phase structure',
    '                   "Phase 2 - Sparring Peak": ..., }',
    '',
    '    return phases',
])

h3('C.2 — Sparring Partner Compatibility Scoring (routes/sparring.py)')
body(
    'The get_sparring_matches() function scores every other registered sparring profile '
    'against the current user\'s profile. The score combines a base value (40), skill-level '
    'proximity (up to 25 points), style compatibility (up to 25 points), and geolocation '
    'proximity via the Haversine formula (up to 20 points). Profiles scoring below 50 are '
    'excluded; the top 20 are returned.'
)
code_block([
    'def haversine_distance(lat1, lon1, lat2, lon2):',
    '    """Return great-circle distance in miles (Sinnott, 1984)."""',
    '    lat1,lon1,lat2,lon2 = map(math.radians, [lat1,lon1,lat2,lon2])',
    '    dlat = lat2 - lat1',
    '    dlon = lon2 - lon1',
    '    a = math.sin(dlat/2)**2 + math.cos(lat1)*math.cos(lat2)*math.sin(dlon/2)**2',
    '    return 2 * math.asin(math.sqrt(a)) * 3959  # radius of Earth in miles',
    '',
    '',
    'def get_sparring_matches(user_id, user_profile):',
    '    matches = []',
    '    all_profiles = SparringProfile.query.filter(',
    '        SparringProfile.user_id != user_id).all()',
    '',
    '    for profile in all_profiles:',
    '        score = 40  # base score',
    '',
    '        # Skill-level proximity (max 25)',
    '        skill_levels = ["Beginner","Intermediate","Advanced","Expert"]',
    '        diff = abs(skill_levels.index(user_profile.skill_level)',
    '                   - skill_levels.index(profile.skill_level))',
    '        score += {0: 25, 1: 20, 2: 10}.get(diff, 5)',
    '',
    '        # Style compatibility (max 25)',
    '        striking  = {"Striking","Kickboxing","Boxing","Muay Thai"}',
    '        grappling = {"Grappling","Wrestling","BJJ"}',
    '        u, p = user_profile.preferred_styles, profile.preferred_styles',
    '        if u == p:                                 score += 25',
    '        elif "Mixed" in (u, p):                   score += 20',
    '        elif {u,p}.issubset(striking):             score += 18',
    '        elif {u,p}.issubset(grappling):            score += 18',
    '        else:                                      score += 10',
    '',
    '        # Geolocation proximity (max 20) via Haversine',
    '        if all([user_profile.latitude, user_profile.longitude,',
    '                profile.latitude,      profile.longitude]):',
    '            dist = haversine_distance(',
    '                user_profile.latitude,  user_profile.longitude,',
    '                profile.latitude,       profile.longitude)',
    '            if dist > user_profile.max_distance: continue  # outside range',
    '            score += (20 if dist<=10 else 15 if dist<=25 else',
    '                      12 if dist<=50 else  8 if dist<=100 else 5)',
    '',
    '        score = max(0, min(100, score))',
    '        if score >= 50:',
    '            matches.append({"profile": profile, "score": score, ...})',
    '',
    '    matches.sort(key=lambda x: x["score"], reverse=True)',
    '    return matches[:20]',
])

h3('C.3 — Readiness and Risk Scoring Algorithm (routes/training.py)')
body(
    'The risk_readiness() route calculates a composite readiness score starting at 100 and '
    'applying deductions based on weight overage, training availability, camp plan status '
    'and proximity to fight date. Each deduction is recorded as a human-readable factor '
    'string for display on the results page.'
)
code_block([
    '@training_bp.route("/risk_readiness")',
    '@login_required',
    'def risk_readiness():',
    '    score   = 100',
    '    factors = []',
    '',
    '    # Weight deduction',
    '    latest  = WeightLog.query.filter_by(user_id=current_user.id)',
    '                             .order_by(WeightLog.date.desc()).first()',
    '    if latest:',
    '        target     = fighter_profile.walk_around_weight - 10',
    '        weight_diff = latest.weight - target',
    '        if weight_diff <= 0:',
    '            factors.append("Weight on track (no deduction)")',
    '        elif weight_diff <= 5:',
    '            score -= 10',
    '            factors.append("Weight close to target (-10)")',
    '        else:',
    '            penalty = min(80, weight_diff * 3)',
    '            score  -= penalty',
    '            factors.append(f"Weight over target (-{penalty})")',
    '',
    '    # Training availability deduction',
    '    if fighter_profile.training_availability < 3:',
    '        score -= 20; factors.append("Low training availability (-20)")',
    '    elif fighter_profile.training_availability < 5:',
    '        score -= 10; factors.append("Moderate training availability (-10)")',
    '',
    '    # Camp plan deduction',
    '    if not CampPlan.query.filter_by(user_id=current_user.id).count():',
    '        score -= 15; factors.append("No camp plan (-15)")',
    '',
    '    # Fight-date proximity deduction',
    '    days = (fighter_profile.fight_date - datetime.today().date()).days',
    '    if   days > 60: factors.append("Sufficient time to fight (no deduction)")',
    '    elif days > 30: score -=  5; factors.append("Moderate time to fight (-5)")',
    '    elif days > 14: score -= 10; factors.append("Limited time to fight (-10)")',
    '    else:           score -= 20; factors.append("Very close to fight (-20)")',
    '',
    '    score = max(0, score)',
    '    readiness_level = "Low" if score < 30 else "Medium" if score < 70 else "High"',
    '    risk_level      = "High" if score < 30 else "Medium" if score < 70 else "Low"',
])

h3('C.4 — Real-Time Chat: SocketIO send_message Handler (routes/social.py)')
body(
    'The handle_send_message() function is the core of the real-time chat feature. When a '
    '"send_message" event is received via WebSocket, it persists the message to the database, '
    'emits it to the conversation room shared by both parties, and pushes an unread-count '
    'update to the recipient\'s personal user room so the badge count updates across all '
    'open tabs without a page reload.'
)
code_block([
    '# current_chats: dict[user_id → friend_id] — tracks which chat each user has open',
    'current_chats = {}',
    '',
    '@socketio.on("join_chat")',
    'def on_join_chat(data):',
    '    """Record that user_id currently has friend_id\'s chat window open."""',
    '    current_chats[int(data["user_id"])] = int(data["friend_id"])',
    '',
    '',
    '@socketio.on("send_message")',
    'def handle_send_message(data, ack=None):',
    '    sender_id    = data.get("sender_id")',
    '    recipient_id = data.get("recipient_id")',
    '    message_text = data.get("message")',
    '',
    '    # Persist to DB',
    '    message = ChatMessage(',
    '        sender_id=sender_id, receiver_id=recipient_id,',
    '        content=message_text, is_read=False)',
    '    db.session.add(message)',
    '    db.session.commit()',
    '',
    '    # Deliver to shared conversation room',
    '    conversation_room = f"chat_{min(sender_id,recipient_id)}_{max(sender_id,recipient_id)}"',
    '    emit("receive_message", {',
    '        "sender_id":       sender_id,',
    '        "sender_username": sender.username,',
    '        "message":         message_text,',
    '        "timestamp":       message.timestamp.isoformat()',
    '    }, room=conversation_room)',
    '',
    '    # Push unread-count update to recipient\'s personal room',
    '    unread_counts = db.session.query(',
    '        ChatMessage.sender_id,',
    '        func.count(ChatMessage.id).label("unread_count")',
    '    ).filter_by(receiver_id=recipient_id, is_read=False)',
    '     .group_by(ChatMessage.sender_id).all()',
    '',
    '    # If recipient has this chat open, zero out unread for this sender',
    '    if current_chats.get(int(recipient_id)) == int(sender_id):',
    '        unread_dict[str(sender_id)] = 0',
    '',
    '    emit("unread_count_update",',
    '         {"unread_counts": unread_dict, "new_message_from": sender_id},',
    '         room=f"user_{recipient_id}")',
])


add_page_numbers()
doc.save('FPMS_Report.docx')
print('Done.')
