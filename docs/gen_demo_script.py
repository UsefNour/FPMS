"""Generate FPMS 20-minute video demo script as a Word document."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

NAVY  = RGBColor(0x1F, 0x49, 0x7D)
BLUE  = RGBColor(0x2E, 0x75, 0xB6)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x00, 0x70, 0x00)
AMBER = RGBColor(0xC5, 0x5A, 0x11)

def add_hrule(color='1F497D'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot); pPr.append(pb)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY
    add_hrule()

def h2(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = color

def body(text, size=11, color=BLACK, italic=False, space_before=0, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.italic = italic

def action(text):
    """Blue bold instruction box — what to DO on screen."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    rb = p.add_run('ACTION: ')
    rb.bold = True; rb.font.size = Pt(11); rb.font.color.rgb = BLUE
    rt = p.add_run(text)
    rt.font.size = Pt(11); rt.font.color.rgb = BLUE

def script(text):
    """Green italic spoken words."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    rb = p.add_run('SAY: ')
    rb.bold = True; rb.font.size = Pt(11); rb.font.color.rgb = GREEN
    rt = p.add_run(f'"{text}"')
    rt.font.size = Pt(11); rt.font.color.rgb = GREEN; rt.italic = True

def tip(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.size = Pt(11); r.font.color.rgb = BLACK

def shade_cell(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def cell_set(cell, text, bold=False, size=11, color=BLACK, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    r.font.color.rgb = color; r.italic = italic

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run('Fighter Performance Management System')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run('20-Minute Video Demo — Script & Instructions')
r2.bold = True; r2.font.size = Pt(16); r2.font.color.rgb = BLUE

add_hrule()

for line, sz, col in [
    ('Youssef Nour  ·  Student ID: 23019868',         12, GREY),
    ('UFCFFF-30-3  —  Software Development Project',  11, GREY),
    ('UWE Bristol  ·  Faculty of Environment and Technology', 11, GREY),
    ('Supervisor: Steve Battle',                       11, GREY),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(line); r.font.size = Pt(sz); r.font.color.rgb = col

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: SETUP
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 1 — Setup Before Recording')

body('Complete all of the following before pressing record:', space_after=6)

setup_items = [
    ('Start the Flask app',         'Open a terminal in the project folder and run:  python app.py\nNavigate to http://127.0.0.1:5000 in your browser.'),
    ('Prepare two accounts',        'Have one admin account and one regular user account already registered.\nLog into the regular account as your primary account for the demo.'),
    ('Populate the fighter DB',     'Run python populate_fighters.py so the fighter database has data to show.'),
    ('Open a second browser',       'Open an Incognito / Private window and log into the second account.\nYou will need this for the friends and chat sections.'),
    ('Recording software',          'Use OBS Studio, Loom, or Windows Xbox Game Bar (Win + G).\nRecord at 1080p with your microphone enabled.'),
    ('Browser',                     'Set browser to fullscreen (F11). Increase zoom to 110% so text is readable on video.'),
    ('Clear browser history',       'No personal data or previous sessions should be visible on screen.'),
]

t = doc.add_table(rows=len(setup_items), cols=2)
t.style = 'Table Grid'
for i, (label, detail) in enumerate(setup_items):
    shade_cell(t.rows[i].cells[0], 'D9E1F2')
    cell_set(t.rows[i].cells[0], label, bold=True, color=NAVY)
    cell_set(t.rows[i].cells[1], detail, size=10.5)
    t.rows[i].cells[0].width = Cm(4.5)
    t.rows[i].cells[1].width = Cm(12.0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: TIMING OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 2 — Timing Overview')

timing = [
    ('0:00 – 1:30',   'Introduction',                    '1.5 min'),
    ('1:30 – 3:00',   'Registration & Login',            '1.5 min'),
    ('3:00 – 4:30',   'Dashboard',                       '1.5 min'),
    ('4:30 – 6:00',   'Fighter Profile',                 '1.5 min'),
    ('6:00 – 8:30',   'Camp Planner',                    '2.5 min'),
    ('8:30 – 10:30',  'Game Plan Builder',               '2.0 min'),
    ('10:30 – 12:00', 'Weight Tracker & Readiness Score','1.5 min'),
    ('12:00 – 13:00', 'Fighter Database',                '1.0 min'),
    ('13:00 – 14:30', 'Friends & Social',                '1.5 min'),
    ('14:30 – 16:00', 'Real-Time Chat',                  '1.5 min'),
    ('16:00 – 17:30', 'Sparring Profile & Partner Finder','1.5 min'),
    ('17:30 – 19:00', 'Events & Admin Panel',            '1.5 min'),
    ('19:00 – 20:00', 'Closing Summary',                 '1.0 min'),
]

tt = doc.add_table(rows=len(timing)+1, cols=3)
tt.style = 'Table Grid'
for col, hdr in enumerate(['Timestamp', 'Section', 'Duration']):
    shade_cell(tt.rows[0].cells[col], '1F497D')
    p = tt.rows[0].cells[col].paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(hdr); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

for i, (ts, sec, dur) in enumerate(timing):
    shade = 'EAF1F9' if i % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([ts, sec, dur]):
        shade_cell(tt.rows[i+1].cells[j], shade)
        cell_set(tt.rows[i+1].cells[j], val,
                 bold=(j==1), color=NAVY if j==1 else BLACK)
    tt.rows[i+1].cells[0].width = Cm(3.5)
    tt.rows[i+1].cells[1].width = Cm(10.0)
    tt.rows[i+1].cells[2].width = Cm(3.0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: FULL SCRIPT
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 3 — Full Demo Script')

body('Blue ACTION lines tell you what to do on screen.\nGreen SAY lines are your spoken words — read them naturally, do not rush.',
     color=GREY, italic=True, space_after=8)

# ── 1. Introduction ───────────────────────────────────────────────────────────
h2('[0:00 – 1:30]  Introduction')
action('Keep the browser on the homepage or the login page.')
script(
    'Hi, my name is Youssef Nour, student number 23019868, and this is my final year project '
    'demo for the module UFCFFF-30-3 at UWE Bristol. The system is called FPMS — the Fighter '
    'Performance Management System. It is a full-stack web application built with Flask, '
    'designed to give combat sports athletes a single platform to manage their entire fight camp. '
    'I will be walking through every feature of the system over the next 20 minutes.'
)

# ── 2. Registration & Login ───────────────────────────────────────────────────
h2('[1:30 – 3:00]  Registration & Login')
action('Navigate to /signup.')
script(
    'Starting with authentication. I will create a new account. The system validates all fields — '
    'unique username, unique email, and password confirmation. Passwords are stored as bcrypt '
    'hashes, never in plain text. The signup form also carries CSRF protection via Flask-WTF '
    'to prevent cross-site request forgery attacks.'
)
action('Fill in the registration form and submit. Then navigate to /login and log in.')
script(
    'On login, Flask-Login creates a secure server-side session and redirects the user to '
    'their personalised dashboard.'
)

# ── 3. Dashboard ─────────────────────────────────────────────────────────────
h2('[3:00 – 4:30]  Dashboard')
action('Show the main dashboard page.')
script(
    'The dashboard is the central hub of FPMS. It shows a live fight countdown based on the '
    'user\'s saved fight date, their current weight status against their target weight class, '
    'which phase of their training camp they are in, and a composite readiness score out of 100. '
    'If any critical thresholds are breached — weight too high, training days low, or no camp '
    'plan set — the system raises warnings here. Everything updates dynamically as the user '
    'logs new data.'
)

# ── 4. Fighter Profile ────────────────────────────────────────────────────────
h2('[4:30 – 6:00]  Fighter Profile')
action('Navigate to /profile.')
script(
    'The fighter profile is where the user sets their core data — weight class, fight date, '
    'fighting style, and their strategic angle: whether they want to take the fight to the ground, '
    'keep it standing, or work from the clinch. This data drives both the camp planner and the '
    'game plan generator, so it is important to fill it in accurately.'
)
action('Fill in the profile with a fight date roughly 10 weeks away and save it.')
script('I\'ll save that and we can see the dashboard countdown update.')

# ── 5. Camp Planner ───────────────────────────────────────────────────────────
h2('[6:00 – 8:30]  Camp Planner')
action('Navigate to /camp_planner.')
script(
    'This is one of the standout features of FPMS. The user enters their opponent\'s name and '
    'a free-text description of the opponent\'s strengths and weaknesses. The system does keyword '
    'analysis on that description — scanning for terms like wrestler, southpaw, pressure fighter, '
    'submission — and generates a periodised training plan tailored to that specific opponent.'
)
action('Fill in the opponent name and a description, then submit.')
script(
    'The plan automatically adapts its phase count based on how many weeks remain until the '
    'fight — 3, 4, or 5 phases. Each phase has a specific training focus, a set of drills, '
    'and a weekly structure. You can generate multiple camp plans for different opponents.'
)
action('Navigate to /view_camp_plans to show the saved plan.')
script('All saved plans are listed here and can be reviewed at any time.')

# ── 6. Game Plan Builder ──────────────────────────────────────────────────────
h2('[8:30 – 10:30]  Game Plan Builder')
action('Navigate to /game_plan.')
script(
    'The game plan builder is different to the camp planner. Rather than periodised training, '
    'it produces a round-by-round tactical blueprint for fight night itself. It uses a 15-branch '
    'decision tree that analyses the opponent profile — preferred range, grappling ability, '
    'striking style — and outputs specific tactical notes for each round: what range to fight at, '
    'techniques to drill, things to avoid, and per-round objectives.'
)
action('Fill in an opponent and generate the game plan. Scroll through the output.')
script(
    'This is saved alongside the camp plan and can be reviewed before fight night. '
    'Together, the camp plan and game plan give the fighter a complete strategic picture '
    'of their preparation.'
)

# ── 7. Weight Tracker & Readiness ────────────────────────────────────────────
h2('[10:30 – 12:00]  Weight Tracker & Readiness Score')
action('Navigate to /weight_tracker.')
script(
    'The weight tracker allows the user to log their weight daily. It displays a 7-day trend '
    'so they can see whether their cut is on track relative to their weight class limit. '
    'The system flags on the dashboard if the gap is becoming dangerous.'
)
action('Log a few weight entries and show the trend chart.')
action('Navigate to /risk_readiness.')
script(
    'The risk and readiness page breaks down every component of the readiness score individually — '
    'weight deviation, training days logged, camp plan presence, and proximity to fight date. '
    'The composite score gives the fighter an at-a-glance picture of how prepared they are '
    'on any given day.'
)

# ── 8. Fighter Database ───────────────────────────────────────────────────────
h2('[12:00 – 13:00]  Fighter Database')
action('Navigate to /fighters.')
script(
    'FPMS includes a built-in fighter database, pre-populated with real combat sports athletes. '
    'Users can browse fighters, filter by weight class and fighting style, and view detailed '
    'profiles including strengths, weaknesses, and preferred range. This is designed to support '
    'opponent research, feeding directly into the camp and game plan tools.'
)
action('Click on a fighter to show their detail page.')

# ── 9. Friends & Social ───────────────────────────────────────────────────────
h2('[13:00 – 14:30]  Friends & Social')
action('Switch to the second browser window logged into the second account. Navigate to /friends.')
script(
    'FPMS has a social layer for connecting training partners and athletes. Users can search '
    'for other registered athletes and send friend requests. The receiving user sees the '
    'request and can accept or decline it.'
)
action('Send a friend request from account 2 to account 1. Switch windows and accept it.')
script(
    'Once connected, users can view each other\'s public fighter profiles. This is useful for '
    'training partners who want to compare preparation and weight class targets.'
)

# ── 10. Real-Time Chat ────────────────────────────────────────────────────────
h2('[14:30 – 16:00]  Real-Time Chat')
action('Open the chat between the two accounts. Place both windows side by side if possible.')
script(
    'The chat system is built on WebSocket connections via Flask-SocketIO. Messages are '
    'delivered in real time without any page refresh. You can see the unread message badge '
    'in the navbar update instantly when a new message arrives. All messages are persisted '
    'to the database, so the conversation history is preserved between sessions.'
)
action('Send a few messages from both accounts and show the badge updating in the navbar.')

# ── 11. Sparring ──────────────────────────────────────────────────────────────
h2('[16:00 – 17:30]  Sparring Profile & Partner Finder')
action('Navigate to /sparring_profile and create a sparring profile.')
script(
    'The sparring module starts with a profile — the user sets their skill level, preferred '
    'styles, and availability. Location is captured to enable distance-based matching.'
)
action('Navigate to /sparring_dashboard.')
script(
    'The partner finder ranks other users by compatibility. It uses the Haversine formula to '
    'calculate real geographic distance between users based on their coordinates, then combines '
    'that with skill level gap and style compatibility to produce a ranked list. Any user '
    'scoring below 50 out of 100 is excluded from results entirely. Users can request a '
    'sparring session directly from a partner\'s card.'
)

# ── 12. Events ────────────────────────────────────────────────────────────────
h2('[17:30 – 19:00]  Events & Admin Panel')
action('Log out and navigate to /events as a guest.')
script(
    'The events module is the only section of FPMS fully accessible without an account. '
    'Guests can browse and filter all upcoming events by date, location, and weight class. '
    'This was a deliberate design decision — lowering the barrier to discovery encourages '
    'more fighters to register on the platform.'
)
action('Log back in and create a new event.')
script(
    'Logged-in users can create events with full details — venue, date, weight classes, '
    'and rules. The organiser manages their listings from the My Events page and can '
    'edit or remove them at any time.'
)
action('Switch to the admin account and navigate to /admin/events.')
script(
    'Admin users have a dedicated panel where they can approve, edit, or remove any event '
    'in the system, giving the platform operator full oversight of all listings.'
)

# ── 13. Closing ───────────────────────────────────────────────────────────────
h2('[19:00 – 20:00]  Closing Summary')
action('Navigate back to the dashboard.')
script(
    'That covers all the core features of FPMS — authentication, fighter profile, camp planning, '
    'game plan generation, weight tracking, readiness scoring, the fighter database, social '
    'connections, real-time chat, sparring partner matching, and event management. '
    'The system was built over four two-to-three week iterations using an Agile approach, '
    'and all 40 functional requirements were verified through manual black-box testing '
    'across 44 routes. Thank you for watching.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: RECORDING TIPS
# ══════════════════════════════════════════════════════════════════════════════
h1('Section 4 — Recording Tips')

tips = [
    'Pause for 2 seconds between sections — gives you clean cut points if you edit the video later.',
    'Zoom in on forms when typing so the text is clearly readable on screen.',
    'Place both browser windows side by side during the chat demo so both accounts are visible simultaneously.',
    'Do not rush — 20 minutes is generous and you have time to breathe between sentences.',
    'If you stumble on a line, pause, take a breath, and re-say it — easy to cut in post-editing.',
    'Keep your mic away from keyboard noise — type slowly during form sections.',
    'Do a 2-minute test recording first to check audio levels and screen visibility.',
    'If using OBS, set your bitrate to at least 4000 kbps for crisp text on screen.',
]

for t in tips:
    tip(t)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(sys.argv[0])),
                   'FPMS_Demo_Script.docx')
doc.save(out)
print(f'Saved: {out}')
